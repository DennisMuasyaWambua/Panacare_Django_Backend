from rest_framework import viewsets, permissions, status, serializers
from rest_framework.response import Response
from rest_framework.decorators import action
from django.utils import timezone
from django.db.models import Q, Avg
from datetime import datetime, timedelta
from django.conf import settings
from .pesapal_client import PesapalClient
from .subscription_utils import SubscriptionManager
from panacare.pagination import CustomPageNumberPagination
from .models import (
    HealthCare, Appointment, Consultation, ConsultationChat, DoctorRating,
    Article, ArticleComment, ArticleCommentLike, PatientDoctorAssignment,
    Package, PatientSubscription, DoctorAvailability, Payment, PatientJournal
    # AppointmentDocument, Resource,
)
from .serializers import (
    HealthCareSerializer, AppointmentSerializer,
    ConsultationSerializer, ConsultationChatSerializer,
    DoctorRatingSerializer, ArticleSerializer, ArticleCommentSerializer, 
    ArticleCommentLikeSerializer, ArticleCommentReplySerializer,
    PackageSerializer, PatientSubscriptionSerializer, DoctorAvailabilitySerializer, PaymentSerializer,
    PatientDoctorAssignmentSerializer, RiskSegmentationSummarySerializer, 
    RiskSegmentationSerializer, PatientRiskListSerializer, PatientJournalSerializer,
  #  AppointmentDocumentSerializer, ResourceSerializer,
)
from doctors.views import IsAdminUser, IsVerifiedUser, IsPatientUser, IsDoctorUser
from users.models import User, Role, Patient
from doctors.models import Doctor
from django.shortcuts import get_object_or_404
from django.http import HttpResponse
from docx import Document
from docx.shared import Inches
import io

from drf_yasg.utils import swagger_auto_schema
from drf_yasg import openapi

# Define the format parameter for Swagger documentation
format_parameter = openapi.Parameter(
    'format', 
    openapi.IN_QUERY, 
    description="Response format. Set to 'fhir' for FHIR-compliant responses", 
    type=openapi.TYPE_STRING,
    required=False,
    enum=['fhir']
)

class HealthCareViewSet(viewsets.ModelViewSet):
    queryset = HealthCare.objects.all()
    serializer_class = HealthCareSerializer
    permission_classes = [permissions.IsAuthenticated]
    
    @swagger_auto_schema(
        operation_description="List all healthcare facilities",
        manual_parameters=[
            format_parameter,
            openapi.Parameter('category', openapi.IN_QUERY, type=openapi.TYPE_STRING, 
                            description="Filter by category (GENERAL, PEDIATRIC, MENTAL, DENTAL, VISION, OTHER)"),
            openapi.Parameter('name', openapi.IN_QUERY, type=openapi.TYPE_STRING, 
                            description="Filter by name (contains search)"),
            openapi.Parameter('active', openapi.IN_QUERY, type=openapi.TYPE_BOOLEAN, 
                            description="Filter by active status")
        ]
    )
    def list(self, request, *args, **kwargs):
        return super().list(request, *args, **kwargs)
        
    @swagger_auto_schema(
        operation_description="Get details of a specific healthcare facility",
        manual_parameters=[format_parameter]
    )
    def retrieve(self, request, *args, **kwargs):
        return super().retrieve(request, *args, **kwargs)
    
    def finalize_response(self, request, response, *args, **kwargs):
        """
        Add CORS headers to all responses
        """
        response = super().finalize_response(request, response, *args, **kwargs)
        response["Access-Control-Allow-Origin"] = "*"
        response["Access-Control-Allow-Methods"] = "GET, POST, PUT, DELETE, OPTIONS"
        response["Access-Control-Allow-Headers"] = "Authorization, Content-Type, X-Requested-With"
        
        # Add FHIR content type header if format is FHIR
        if request.query_params.get('format') == 'fhir':
            response["Content-Type"] = "application/fhir+json"
            
        return response
    
    def get_permissions(self):
        """
        Override to set custom permissions for different actions
        """
        if self.action in ['create', 'update', 'partial_update', 'destroy']:
            permission_classes = [IsAdminUser]
        else:
            permission_classes = [permissions.IsAuthenticated]
        return [permission() for permission in permission_classes]
    
    def get_queryset(self):
        queryset = HealthCare.objects.all()
        category = self.request.query_params.get('category')
        if category:
            queryset = queryset.filter(category=category)
        name = self.request.query_params.get('name')
        if name:
            queryset = queryset.filter(name__icontains=name)
        active = self.request.query_params.get('active')
        if active:
            queryset = queryset.filter(is_active=active.lower() == 'true')
        return queryset
    
    @action(detail=False, methods=['post'], permission_classes=[IsAdminUser])
    def assign_patient_to_doctor(self, request):
        """
        Endpoint for admin to assign a patient to a doctor
        """
        # Get patient and doctor
        patient_id = request.data.get('patient_id')
        doctor_id = request.data.get('doctor_id')
        notes = request.data.get('notes', '')
        
        if not patient_id or not doctor_id:
            return Response({
                'error': 'patient_id and doctor_id are required'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            patient = Patient.objects.get(id=patient_id)
            doctor = Doctor.objects.get(id=doctor_id)
        except (Patient.DoesNotExist, Doctor.DoesNotExist):
            return Response({
                'error': 'Patient or doctor not found'
            }, status=status.HTTP_404_NOT_FOUND)
        
        # Check if assignment already exists
        existing = PatientDoctorAssignment.objects.filter(
            patient=patient, 
            doctor=doctor,
            is_active=True
        ).exists()
        
        if existing:
            return Response({
                'error': 'This patient is already assigned to this doctor'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # Create new assignment
        assignment = PatientDoctorAssignment.objects.create(
            patient=patient,
            doctor=doctor,
            notes=notes
        )
        
        serializer = PatientDoctorAssignmentSerializer(assignment)
        return Response(serializer.data, status=status.HTTP_201_CREATED)
    
    @swagger_auto_schema(
        method='get',
        operation_description="List all patient-doctor assignments",
        manual_parameters=[
            format_parameter,
            openapi.Parameter('doctor_id', openapi.IN_QUERY, type=openapi.TYPE_STRING, 
                            description="Filter by doctor ID"),
            openapi.Parameter('patient_id', openapi.IN_QUERY, type=openapi.TYPE_STRING, 
                            description="Filter by patient ID")
        ]
    )
    @action(detail=False, methods=['get'], permission_classes=[IsAdminUser])
    def list_patient_doctor_assignments(self, request):
        """
        Endpoint for admin to list all patient-doctor assignments
        """
        assignments = PatientDoctorAssignment.objects.filter(is_active=True)
        
        # Optional filtering
        doctor_id = request.query_params.get('doctor_id')
        if doctor_id:
            assignments = assignments.filter(doctor_id=doctor_id)
            
        patient_id = request.query_params.get('patient_id')
        if patient_id:
            assignments = assignments.filter(patient_id=patient_id)
        
        serializer = PatientDoctorAssignmentSerializer(assignments, many=True, context={'request': request})
        
        # Set appropriate content type for FHIR responses
        response = Response(serializer.data)
        if request.query_params.get('format') == 'fhir':
            response["Content-Type"] = "application/fhir+json"
            
        return response
    
    @action(detail=True, methods=['get'], permission_classes=[IsAdminUser])
    def view_assignment(self, request, pk=None):
        """
        Endpoint for admin to view a specific patient-doctor assignment
        """
        assignment = get_object_or_404(PatientDoctorAssignment, pk=pk)
        serializer = PatientDoctorAssignmentSerializer(assignment)
        return Response(serializer.data)
    
    @swagger_auto_schema(
        method='get',
        operation_description="List all patients assigned to the doctor",
        manual_parameters=[
            format_parameter,
            openapi.Parameter('active', openapi.IN_QUERY, type=openapi.TYPE_BOOLEAN, 
                            description="Filter by active status (true/false)")
        ]
    )
    @action(detail=False, methods=['get'], permission_classes=[IsDoctorUser])
    def doctor_assigned_patients(self, request):
        """
        Endpoint for doctors to view all patients assigned to them
        """
        try:
            # Get doctor from user
            doctor = request.user.doctor
            
            # Get all active assignments for this doctor
            assignments = PatientDoctorAssignment.objects.filter(doctor=doctor)
            
            # Optional filtering by active status
            is_active = request.query_params.get('active')
            if is_active is not None:
                assignments = assignments.filter(is_active=is_active.lower() == 'true')
                
            # Get patient objects from assignments
            patient_ids = assignments.values_list('patient_id', flat=True)
            patients = Patient.objects.filter(id__in=patient_ids)
            
            # Use the patient serializer 
            from users.serializers import PatientSerializer
            serializer = PatientSerializer(patients, many=True, context={'request': request})
            
            # Set appropriate content type for FHIR responses
            response = Response(serializer.data)
            if request.query_params.get('format') == 'fhir':
                response["Content-Type"] = "application/fhir+json"
                
            return response
        except Doctor.DoesNotExist:
            return Response({
                'error': 'Doctor profile not found'
            }, status=status.HTTP_404_NOT_FOUND)
    
    @swagger_auto_schema(
        method='get',
        operation_description="Get details of a specific patient assigned to the doctor",
        manual_parameters=[format_parameter]
    )
    @action(detail=True, methods=['get'], permission_classes=[IsDoctorUser])
    def doctor_view_patient(self, request, pk=None):
        """
        Endpoint for doctors to view details of a specific patient assigned to them
        """
        try:
            # Get doctor from user
            doctor = request.user.doctor
            
            # Get the patient
            patient = get_object_or_404(Patient, pk=pk)
            
            # Check if the patient is assigned to this doctor
            assignment_exists = PatientDoctorAssignment.objects.filter(
                doctor=doctor,
                patient=patient,
                is_active=True
            ).exists()
            
            if not assignment_exists:
                return Response({
                    'error': 'This patient is not assigned to you or the assignment is not active'
                }, status=status.HTTP_403_FORBIDDEN)
            
            # Use the patient serializer
            from users.serializers import PatientSerializer
            serializer = PatientSerializer(patient, context={'request': request})
            
            # Set appropriate content type for FHIR responses
            response = Response(serializer.data)
            if request.query_params.get('format') == 'fhir':
                response["Content-Type"] = "application/fhir+json"
                
            return response
        except Doctor.DoesNotExist:
            return Response({
                'error': 'Doctor profile not found'
            }, status=status.HTTP_404_NOT_FOUND)


class IsPatientUser(permissions.BasePermission):
    """
    Permission class to check if the user has patient role
    """
    def has_permission(self, request, view):
        # Check if user is authenticated
        if not request.user.is_authenticated:
            return False
        
        # Check if user has patient role
        return request.user.roles.filter(name='patient').exists()




class IsPatientOrDoctorOrAdmin(permissions.BasePermission):
    """
    Permission class to check if the user has patient, doctor or admin role
    """
    def has_permission(self, request, view):
        # Check if user is authenticated
        if not request.user.is_authenticated:
            return False
        
        # Check if user has patient, doctor or admin role
        return request.user.roles.filter(name__in=['patient', 'doctor', 'admin']).exists()


# class DoctorAvailabilityViewSet(viewsets.ModelViewSet):
#     queryset = DoctorAvailability.objects.all()
#     serializer_class = DoctorAvailabilitySerializer
#     permission_classes = [permissions.IsAuthenticated]
#     
#     def get_permissions(self):
#         """
#         Override to set custom permissions for different actions
#         """
#         if self.action in ['create', 'update', 'partial_update', 'destroy']:
#             permission_classes = [IsDoctorUser | IsAdminUser]
#         else:
#             permission_classes = [permissions.IsAuthenticated]
#         return [permission() for permission in permission_classes]
#     
#     def get_queryset(self):
#         queryset = DoctorAvailability.objects.all()
#         
#         # Optional filtering
#         doctor_id = self.request.query_params.get('doctor_id')
#         if doctor_id:
#             queryset = queryset.filter(doctor_id=doctor_id)
#             
#         is_available = self.request.query_params.get('available')
#         if is_available:
#             queryset = queryset.filter(is_available=is_available.lower() == 'true')
#             
#         day_of_week = self.request.query_params.get('day')
#         if day_of_week:
#             try:
#                 day_of_week = int(day_of_week)
#                 queryset = queryset.filter(day_of_week=day_of_week)
#             except ValueError:
#                 pass
#                 
#         date = self.request.query_params.get('date')
#         if date:
#             try:
#                 specific_date = datetime.strptime(date, '%Y-%m-%d').date()
#                 queryset = queryset.filter(
#                     Q(specific_date=specific_date) | 
#                     (Q(is_recurring=True) & Q(day_of_week=specific_date.weekday()))
#                 )
#             except ValueError:
#                 pass
#                 
#         return queryset
#     
#     @action(detail=False, methods=['get'], permission_classes=[IsDoctorUser])
#     def my_availability(self, request):
#         """
#         Endpoint for doctors to view their own availability
#         """
#         try:
#             doctor = request.user.doctor
#             availabilities = DoctorAvailability.objects.filter(doctor=doctor)
#             serializer = self.get_serializer(availabilities, many=True)
#             return Response(serializer.data)
#         except Doctor.DoesNotExist:
#             return Response({
#                 'error': 'Doctor profile not found'
#             }, status=status.HTTP_404_NOT_FOUND)
# 
# 
class AppointmentViewSet(viewsets.ModelViewSet):
    """
    ViewSet for managing appointments
    
    Status Flow:
    - booked: Initial appointment status
    - arrived: When consultation starts (appointment.status changes to 'arrived')
    - fulfilled: When consultation ends (appointment.status changes to 'fulfilled')
    
    Note: Patients should be able to see all their appointments regardless of status.
    Frontend apps should not filter by status to hide 'arrived' or 'fulfilled' appointments.
    """
    queryset = Appointment.objects.all()
    serializer_class = AppointmentSerializer
    permission_classes = [permissions.IsAuthenticated]
    
    def get_permissions(self):
        """
        Override to set custom permissions for different actions
        """
        if self.action in ['create']:
            permission_classes = [IsPatientUser | IsAdminUser]
        elif self.action in ['update', 'partial_update']:
            permission_classes = [IsPatientOrDoctorOrAdmin]
        elif self.action in ['destroy']:
            permission_classes = [IsAdminUser]
        else:
            permission_classes = [permissions.IsAuthenticated]
        return [permission() for permission in permission_classes]
    
    def get_queryset(self):
        queryset = Appointment.objects.all()
        
        # Filter by role
        if self.request.user.is_authenticated and self.request.user.roles.filter(name='admin').exists():
            # Admin can see all appointments
            pass
        elif self.request.user.is_authenticated and self.request.user.roles.filter(name='doctor').exists():
            # Doctors can only see their own appointments
            try:
                doctor = self.request.user.doctor
                queryset = queryset.filter(doctor=doctor)
            except Doctor.DoesNotExist:
                return Appointment.objects.none()
        elif self.request.user.is_authenticated and self.request.user.roles.filter(name='patient').exists():
            # Patients can only see their own appointments
            try:
                patient = self.request.user.patient
                queryset = queryset.filter(patient=patient)
            except Patient.DoesNotExist:
                return Appointment.objects.none()
        elif self.request.user.is_authenticated and self.request.user.roles.filter(name='community_health_provider').exists():
            # CHPs can only see appointments they created
            try:
                from users.models import CommunityHealthProvider
                chp = self.request.user.community_health_provider
                queryset = queryset.filter(created_by_chp=chp)
            except CommunityHealthProvider.DoesNotExist:
                return Appointment.objects.none()
        else:
            return Appointment.objects.none()
            
        # Optional filtering
        doctor_id = self.request.query_params.get('doctor_id')
        if doctor_id:
            queryset = queryset.filter(doctor_id=doctor_id)
            
        patient_id = self.request.query_params.get('patient_id')
        if patient_id:
            queryset = queryset.filter(patient_id=patient_id)
            
        status_param = self.request.query_params.get('status')
        if status_param:
            queryset = queryset.filter(status=status_param)
            
        date_from = self.request.query_params.get('date_from')
        if date_from:
            try:
                date_from = datetime.strptime(date_from, '%Y-%m-%d').date()
                queryset = queryset.filter(appointment_date__gte=date_from)
            except ValueError:
                pass
                
        date_to = self.request.query_params.get('date_to')
        if date_to:
            try:
                date_to = datetime.strptime(date_to, '%Y-%m-%d').date()
                queryset = queryset.filter(appointment_date__lte=date_to)
            except ValueError:
                pass
                
        appointment_type = self.request.query_params.get('type')
        if appointment_type:
            queryset = queryset.filter(appointment_type=appointment_type)
                
        return queryset
    
    @action(detail=False, methods=['get'], permission_classes=[IsPatientUser])
    def my_appointments(self, request):
        """
        Endpoint for patients to view their own appointments
        """
        try:
            patient = request.user.patient
            appointments = Appointment.objects.filter(patient=patient)
            
            # Optional filtering
            status_param = request.query_params.get('status')
            if status_param:
                appointments = appointments.filter(status=status_param)
                
            date_from = request.query_params.get('date_from')
            if date_from:
                try:
                    date_from = datetime.strptime(date_from, '%Y-%m-%d').date()
                    appointments = appointments.filter(appointment_date__gte=date_from)
                except ValueError:
                    pass
                    
            date_to = request.query_params.get('date_to')
            if date_to:
                try:
                    date_to = datetime.strptime(date_to, '%Y-%m-%d').date()
                    appointments = appointments.filter(appointment_date__lte=date_to)
                except ValueError:
                    pass
                    
            doctor_id = request.query_params.get('doctor_id')
            if doctor_id:
                appointments = appointments.filter(doctor_id=doctor_id)
                
            serializer = self.get_serializer(appointments, many=True)
            return Response(serializer.data)
        except Patient.DoesNotExist:
            return Response({
                'error': 'Patient profile not found'
            }, status=status.HTTP_404_NOT_FOUND)
    
    @action(detail=False, methods=['get'], permission_classes=[IsDoctorUser])
    def doctor_appointments(self, request):
        """
        Endpoint for doctors to view their own appointments
        """
        try:
            doctor = request.user.doctor
            appointments = Appointment.objects.filter(doctor=doctor)
            
            # Optional filtering
            status_param = request.query_params.get('status')
            if status_param:
                appointments = appointments.filter(status=status_param)
                
            date_from = request.query_params.get('date_from')
            if date_from:
                try:
                    date_from = datetime.strptime(date_from, '%Y-%m-%d').date()
                    appointments = appointments.filter(appointment_date__gte=date_from)
                except ValueError:
                    pass
                    
            date_to = request.query_params.get('date_to')
            if date_to:
                try:
                    date_to = datetime.strptime(date_to, '%Y-%m-%d').date()
                    appointments = appointments.filter(appointment_date__lte=date_to)
                except ValueError:
                    pass
                    
            patient_id = request.query_params.get('patient_id')
            if patient_id:
                appointments = appointments.filter(patient_id=patient_id)
                
            serializer = self.get_serializer(appointments, many=True)
            return Response(serializer.data)
        except Doctor.DoesNotExist:
            return Response({
                'error': 'Doctor profile not found'
            }, status=status.HTTP_404_NOT_FOUND)
            
    @action(detail=True, methods=['post'], permission_classes=[IsPatientUser])
    def cancel_appointment(self, request, pk=None):
        """
        Endpoint for patients to cancel their appointments
        """
        appointment = self.get_object()
        
        # Check if patient owns this appointment
        if not appointment.patient.user == request.user:
            return Response({
                'error': 'You can only cancel your own appointments'
            }, status=status.HTTP_403_FORBIDDEN)
            
        # Check if appointment can be cancelled
        if appointment.status in ['fulfilled', 'cancelled', 'noshow']:
            return Response({
                'error': f'Cannot cancel an appointment with status: {appointment.status}'
            }, status=status.HTTP_400_BAD_REQUEST)
            
        # Cancel the appointment
        appointment.status = 'cancelled'
        appointment.save()
        
        serializer = self.get_serializer(appointment)
        return Response(serializer.data)
        
    @swagger_auto_schema(
        operation_description="Reschedule an appointment as a patient",
        request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            required=['appointment_date', 'start_time', 'end_time', 'reschedule_reason'],
            properties={
                'appointment_date': openapi.Schema(type=openapi.TYPE_STRING, format='date', description='New appointment date (YYYY-MM-DD)'),
                'start_time': openapi.Schema(type=openapi.TYPE_STRING, format='time', description='New start time (HH:MM:SS)'),
                'end_time': openapi.Schema(type=openapi.TYPE_STRING, format='time', description='New end time (HH:MM:SS)'),
                'reschedule_reason': openapi.Schema(type=openapi.TYPE_STRING, description='Reason for rescheduling the appointment')
            }
        ),
        responses={
            200: AppointmentSerializer,
            400: openapi.Response("Bad Request", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'error': openapi.Schema(type=openapi.TYPE_STRING, description="Error message")
                }
            )),
            403: openapi.Response("Forbidden", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'error': openapi.Schema(type=openapi.TYPE_STRING, description="Permission error message")
                }
            ))
        }
    )
    @action(detail=True, methods=['post'], permission_classes=[IsPatientUser])
    def patient_reschedule(self, request, pk=None):
        """
        Endpoint for patients to reschedule their appointments
        """
        appointment = self.get_object()
        
        # Check if patient owns this appointment
        if not appointment.patient.user == request.user:
            return Response({
                'error': 'You can only reschedule your own appointments'
            }, status=status.HTTP_403_FORBIDDEN)
            
        # Check if appointment can be rescheduled
        if appointment.status in ['fulfilled', 'cancelled', 'noshow']:
            return Response({
                'error': f'Cannot reschedule an appointment with status: {appointment.status}'
            }, status=status.HTTP_400_BAD_REQUEST)
            
        # Get new appointment details
        appointment_date = request.data.get('appointment_date')
        start_time = request.data.get('start_time')
        end_time = request.data.get('end_time')
        reschedule_reason = request.data.get('reschedule_reason')
        
        if not all([appointment_date, start_time, end_time, reschedule_reason]):
            return Response({
                'error': 'appointment_date, start_time, end_time, and reschedule_reason are required'
            }, status=status.HTTP_400_BAD_REQUEST)
            
        try:
            # Convert to appropriate formats
            from datetime import datetime
            appointment_date = datetime.strptime(appointment_date, '%Y-%m-%d').date()
        except ValueError:
            return Response({
                'error': 'Invalid date format. Use YYYY-MM-DD format.'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # Store the previous values for reference
        previous_date = appointment.appointment_date
        previous_start = appointment.start_time
        previous_end = appointment.end_time
        
        # Update the appointment
        appointment.appointment_date = appointment_date
        appointment.start_time = start_time
        appointment.end_time = end_time
        
        # Add rescheduling reason to notes
        current_notes = appointment.notes or ""
        reschedule_note = f"[Patient Reschedule {datetime.now().strftime('%Y-%m-%d %H:%M')}] "\
                         f"Changed from {previous_date} {previous_start}-{previous_end} to "\
                         f"{appointment_date} {start_time}-{end_time}. "\
                         f"Reason: {reschedule_reason}"
        
        if current_notes:
            appointment.notes = f"{current_notes}\n\n{reschedule_note}"
        else:
            appointment.notes = reschedule_note
        
        # Update status to rescheduled/scheduled
        appointment.status = 'scheduled'
        appointment.save()
        
        serializer = self.get_serializer(appointment)
        return Response(serializer.data)
    
    @action(detail=True, methods=['post'], permission_classes=[IsDoctorUser])
    def update_consultation_details(self, request, pk=None):
        """
        Endpoint for doctors to update consultation details (diagnosis, treatment, etc.)
        """
        appointment = self.get_object()
        
        # Check if doctor owns this appointment
        if not appointment.doctor.user == request.user:
            return Response({
                'error': 'You can only update your own appointments'
            }, status=status.HTTP_403_FORBIDDEN)
            
        # Update the appointment
        diagnosis = request.data.get('diagnosis')
        if diagnosis is not None:
            appointment.diagnosis = diagnosis
            
        treatment = request.data.get('treatment')
        if treatment is not None:
            appointment.treatment = treatment
            
        notes = request.data.get('notes')
        if notes is not None:
            appointment.notes = notes
            
        risk_level = request.data.get('risk_level')
        if risk_level is not None:
            appointment.risk_level = risk_level
            
        # Change status to fulfilled if requested
        status_param = request.data.get('status')
        if status_param:
            appointment.status = status_param
            
        appointment.save()
        
        serializer = self.get_serializer(appointment)
        return Response(serializer.data)
        
    @swagger_auto_schema(
        operation_description="Reschedule an appointment as a doctor",
        request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            required=['appointment_date', 'start_time', 'end_time', 'reschedule_reason'],
            properties={
                'appointment_date': openapi.Schema(type=openapi.TYPE_STRING, format='date', description='New appointment date (YYYY-MM-DD)'),
                'start_time': openapi.Schema(type=openapi.TYPE_STRING, format='time', description='New start time (HH:MM:SS)'),
                'end_time': openapi.Schema(type=openapi.TYPE_STRING, format='time', description='New end time (HH:MM:SS)'),
                'reschedule_reason': openapi.Schema(type=openapi.TYPE_STRING, description='Reason for rescheduling the appointment')
            }
        )
    )
    @action(detail=True, methods=['post'], permission_classes=[IsDoctorUser])
    def doctor_reschedule(self, request, pk=None):
        """
        Endpoint for doctors to reschedule appointments
        """
        appointment = self.get_object()
        
        # Check if doctor owns this appointment
        if not appointment.doctor.user == request.user:
            return Response({
                'error': 'You can only reschedule your own appointments'
            }, status=status.HTTP_403_FORBIDDEN)
            
        # Check if appointment can be rescheduled
        if appointment.status in ['fulfilled', 'cancelled', 'noshow']:
            return Response({
                'error': f'Cannot reschedule an appointment with status: {appointment.status}'
            }, status=status.HTTP_400_BAD_REQUEST)
            
        # Get new appointment details
        appointment_date = request.data.get('appointment_date')
        start_time = request.data.get('start_time')
        end_time = request.data.get('end_time')
        reschedule_reason = request.data.get('reschedule_reason')
        
        if not all([appointment_date, start_time, end_time, reschedule_reason]):
            return Response({
                'error': 'appointment_date, start_time, end_time, and reschedule_reason are required'
            }, status=status.HTTP_400_BAD_REQUEST)
            
        try:
            # Convert to appropriate formats
            from datetime import datetime
            appointment_date = datetime.strptime(appointment_date, '%Y-%m-%d').date()
        except ValueError:
            return Response({
                'error': 'Invalid date format. Use YYYY-MM-DD format.'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # Store the previous values for reference
        previous_date = appointment.appointment_date
        previous_start = appointment.start_time
        previous_end = appointment.end_time
        
        # Update the appointment
        appointment.appointment_date = appointment_date
        appointment.start_time = start_time
        appointment.end_time = end_time
        
        # Add rescheduling reason to notes
        current_notes = appointment.notes or ""
        reschedule_note = f"[Doctor Reschedule {datetime.now().strftime('%Y-%m-%d %H:%M')}] "\
                         f"Changed from {previous_date} {previous_start}-{previous_end} to "\
                         f"{appointment_date} {start_time}-{end_time}. "\
                         f"Reason: {reschedule_reason}"
        
        if current_notes:
            appointment.notes = f"{current_notes}\n\n{reschedule_note}"
        else:
            appointment.notes = reschedule_note
        
        # Update status to rescheduled/scheduled
        appointment.status = 'scheduled'
        appointment.save()
        
        serializer = self.get_serializer(appointment)
        return Response(serializer.data)
    
    @swagger_auto_schema(
        operation_description="Export single appointment as PDF",
        responses={
            200: openapi.Response(
                description="PDF file",
                schema=openapi.Schema(type=openapi.TYPE_FILE)
            ),
            403: openapi.Response("Forbidden", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'error': openapi.Schema(type=openapi.TYPE_STRING, description="Permission error message")
                }
            )),
            404: openapi.Response("Not Found", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'error': openapi.Schema(type=openapi.TYPE_STRING, description="Appointment not found")
                }
            ))
        }
    )
    @action(detail=True, methods=['get'])
    def export_pdf(self, request, pk=None):
        """
        Export a single appointment as PDF
        """
        from reportlab.lib.pagesizes import A4, letter
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.lib.units import inch
        import io
        
        appointment = self.get_object()
        
        # Check permissions - users can only export their own appointments or admin can export any
        if not (self.request.user.is_authenticated and self.request.user.roles.filter(name='admin').exists()):
            if hasattr(self.request.user, 'patient') and appointment.patient.user != self.request.user:
                return Response({
                    'error': 'You can only export your own appointments'
                }, status=status.HTTP_403_FORBIDDEN)
            elif hasattr(self.request.user, 'doctor') and appointment.doctor.user != self.request.user:
                return Response({
                    'error': 'You can only export your own appointments'
                }, status=status.HTTP_403_FORBIDDEN)
        
        # Create PDF buffer
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
        story = []
        
        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=colors.darkblue
        )
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            textColor=colors.darkblue
        )
        normal_style = styles['Normal']
        
        # Title
        story.append(Paragraph("Appointment Report", title_style))
        story.append(Spacer(1, 20))
        
        # Basic Information Table
        basic_data = [
            ['Appointment ID:', str(appointment.id)],
            ['Date:', appointment.appointment_date.strftime('%B %d, %Y')],
            ['Time:', f"{appointment.start_time.strftime('%I:%M %p')} - {appointment.end_time.strftime('%I:%M %p')}"],
            ['Status:', appointment.get_status_display()],
            ['Type:', appointment.get_appointment_type_display()],
            ['Risk Level:', appointment.get_risk_level_display() if appointment.risk_level else 'Not specified'],
        ]
        
        basic_table = Table(basic_data, colWidths=[2*inch, 4*inch])
        basic_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
            ('ALIGN', (1, 0), (1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
        ]))
        
        story.append(basic_table)
        story.append(Spacer(1, 20))
        
        # Patient Information
        story.append(Paragraph("Patient Information", heading_style))
        patient_data = [
            ['Name:', f"{appointment.patient.user.first_name} {appointment.patient.user.last_name}"],
            ['Email:', appointment.patient.user.email],
            ['Phone:', appointment.patient.phone_number or 'Not provided'],
            ['Date of Birth:', appointment.patient.date_of_birth.strftime('%B %d, %Y') if appointment.patient.date_of_birth else 'Not provided'],
        ]
        
        patient_table = Table(patient_data, colWidths=[2*inch, 4*inch])
        patient_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
            ('ALIGN', (1, 0), (1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
        ]))
        
        story.append(patient_table)
        story.append(Spacer(1, 20))
        
        # Doctor Information
        story.append(Paragraph("Doctor Information", heading_style))
        doctor_data = [
            ['Name:', f"Dr. {appointment.doctor.user.first_name} {appointment.doctor.user.last_name}"],
            ['Email:', appointment.doctor.user.email],
            ['Specialization:', appointment.doctor.specialization or 'Not specified'],
            ['License Number:', appointment.doctor.license_number or 'Not specified'],
        ]
        
        doctor_table = Table(doctor_data, colWidths=[2*inch, 4*inch])
        doctor_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
            ('ALIGN', (1, 0), (1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
        ]))
        
        story.append(doctor_table)
        story.append(Spacer(1, 20))
        
        # Healthcare Facility
        if appointment.healthcare_facility:
            story.append(Paragraph("Healthcare Facility", heading_style))
            facility_data = [
                ['Name:', appointment.healthcare_facility.name],
                ['Address:', appointment.healthcare_facility.address or 'Not specified'],
                ['Phone:', appointment.healthcare_facility.phone_number or 'Not specified'],
            ]
            
            facility_table = Table(facility_data, colWidths=[2*inch, 4*inch])
            facility_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
                ('ALIGN', (1, 0), (1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ]))
            
            story.append(facility_table)
            story.append(Spacer(1, 20))
        
        # Appointment Details
        if appointment.reason or appointment.diagnosis or appointment.treatment or appointment.notes:
            story.append(Paragraph("Appointment Details", heading_style))
            
            if appointment.reason:
                story.append(Paragraph("<b>Reason for Visit:</b>", normal_style))
                story.append(Paragraph(appointment.reason, normal_style))
                story.append(Spacer(1, 12))
            
            if appointment.diagnosis:
                story.append(Paragraph("<b>Diagnosis:</b>", normal_style))
                story.append(Paragraph(appointment.diagnosis, normal_style))
                story.append(Spacer(1, 12))
            
            if appointment.treatment:
                story.append(Paragraph("<b>Treatment:</b>", normal_style))
                story.append(Paragraph(appointment.treatment, normal_style))
                story.append(Spacer(1, 12))
            
            if appointment.notes:
                story.append(Paragraph("<b>Notes:</b>", normal_style))
                story.append(Paragraph(appointment.notes, normal_style))
                story.append(Spacer(1, 12))
        
        # Footer
        story.append(Spacer(1, 30))
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=8,
            alignment=TA_CENTER,
            textColor=colors.grey
        )
        story.append(Paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", footer_style))
        story.append(Paragraph("Panacare Healthcare System", footer_style))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        
        # Create response
        response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
        filename = f"appointment_{appointment.id}_{appointment.appointment_date.strftime('%Y%m%d')}.pdf"
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        buffer.close()
        return response
    
    @swagger_auto_schema(
        operation_description="Advanced appointment search with comprehensive filters",
        manual_parameters=[
            openapi.Parameter('search', openapi.IN_QUERY, type=openapi.TYPE_STRING, description='Search in patient names, doctor names, healthcare facility names'),
            openapi.Parameter('patient_name', openapi.IN_QUERY, type=openapi.TYPE_STRING, description='Filter by patient name'),
            openapi.Parameter('doctor_name', openapi.IN_QUERY, type=openapi.TYPE_STRING, description='Filter by doctor name'),
            openapi.Parameter('healthcare_facility', openapi.IN_QUERY, type=openapi.TYPE_STRING, description='Filter by healthcare facility name'),
            openapi.Parameter('status', openapi.IN_QUERY, type=openapi.TYPE_STRING, description='Filter by appointment status'),
            openapi.Parameter('appointment_type', openapi.IN_QUERY, type=openapi.TYPE_STRING, description='Filter by appointment type'),
            openapi.Parameter('risk_level', openapi.IN_QUERY, type=openapi.TYPE_STRING, description='Filter by risk level'),
            openapi.Parameter('date_from', openapi.IN_QUERY, type=openapi.TYPE_STRING, format='date', description='Filter appointments from date (YYYY-MM-DD)'),
            openapi.Parameter('date_to', openapi.IN_QUERY, type=openapi.TYPE_STRING, format='date', description='Filter appointments to date (YYYY-MM-DD)'),
            openapi.Parameter('time_from', openapi.IN_QUERY, type=openapi.TYPE_STRING, format='time', description='Filter appointments from time (HH:MM)'),
            openapi.Parameter('time_to', openapi.IN_QUERY, type=openapi.TYPE_STRING, format='time', description='Filter appointments to time (HH:MM)'),
            openapi.Parameter('patient_id', openapi.IN_QUERY, type=openapi.TYPE_STRING, description='Filter by patient ID'),
            openapi.Parameter('doctor_id', openapi.IN_QUERY, type=openapi.TYPE_STRING, description='Filter by doctor ID'),
            openapi.Parameter('subscription_package', openapi.IN_QUERY, type=openapi.TYPE_STRING, description='Filter by patient subscription package'),
            openapi.Parameter('has_diagnosis', openapi.IN_QUERY, type=openapi.TYPE_BOOLEAN, description='Filter appointments with diagnosis'),
            openapi.Parameter('has_treatment', openapi.IN_QUERY, type=openapi.TYPE_BOOLEAN, description='Filter appointments with treatment'),
            openapi.Parameter('ordering', openapi.IN_QUERY, type=openapi.TYPE_STRING, description='Order by: appointment_date, start_time, patient_name, doctor_name, status (add - for descending)'),
        ],
        responses={
            200: AppointmentSerializer(many=True),
            400: openapi.Response("Bad Request", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'error': openapi.Schema(type=openapi.TYPE_STRING, description="Error message")
                }
            ))
        }
    )
    @action(detail=False, methods=['get'])
    def search_appointments(self, request):
        """
        Advanced search endpoint for appointments with comprehensive filtering
        """
        queryset = self.get_queryset()
        
        # Text search across multiple fields
        search_query = request.query_params.get('search')
        if search_query:
            from django.db.models import Q
            queryset = queryset.filter(
                Q(patient__user__first_name__icontains=search_query) |
                Q(patient__user__last_name__icontains=search_query) |
                Q(doctor__user__first_name__icontains=search_query) |
                Q(doctor__user__last_name__icontains=search_query) |
                Q(healthcare_facility__name__icontains=search_query) |
                Q(reason__icontains=search_query) |
                Q(diagnosis__icontains=search_query) |
                Q(treatment__icontains=search_query) |
                Q(notes__icontains=search_query)
            )
        
        # Specific field filters
        patient_name = request.query_params.get('patient_name')
        if patient_name:
            queryset = queryset.filter(
                Q(patient__user__first_name__icontains=patient_name) |
                Q(patient__user__last_name__icontains=patient_name)
            )
        
        doctor_name = request.query_params.get('doctor_name')
        if doctor_name:
            queryset = queryset.filter(
                Q(doctor__user__first_name__icontains=doctor_name) |
                Q(doctor__user__last_name__icontains=doctor_name)
            )
        
        healthcare_facility = request.query_params.get('healthcare_facility')
        if healthcare_facility:
            queryset = queryset.filter(healthcare_facility__name__icontains=healthcare_facility)
        
        # Status filter
        status_filter = request.query_params.get('status')
        if status_filter:
            queryset = queryset.filter(status=status_filter)
        
        # Appointment type filter
        appointment_type = request.query_params.get('appointment_type')
        if appointment_type:
            queryset = queryset.filter(appointment_type=appointment_type)
        
        # Risk level filter
        risk_level = request.query_params.get('risk_level')
        if risk_level:
            queryset = queryset.filter(risk_level=risk_level)
        
        # Date range filters
        date_from = request.query_params.get('date_from')
        if date_from:
            try:
                date_from = datetime.strptime(date_from, '%Y-%m-%d').date()
                queryset = queryset.filter(appointment_date__gte=date_from)
            except ValueError:
                return Response({
                    'error': 'Invalid date_from format. Use YYYY-MM-DD format.'
                }, status=status.HTTP_400_BAD_REQUEST)
        
        date_to = request.query_params.get('date_to')
        if date_to:
            try:
                date_to = datetime.strptime(date_to, '%Y-%m-%d').date()
                queryset = queryset.filter(appointment_date__lte=date_to)
            except ValueError:
                return Response({
                    'error': 'Invalid date_to format. Use YYYY-MM-DD format.'
                }, status=status.HTTP_400_BAD_REQUEST)
        
        # Time range filters
        time_from = request.query_params.get('time_from')
        if time_from:
            try:
                from datetime import time
                time_from = datetime.strptime(time_from, '%H:%M').time()
                queryset = queryset.filter(start_time__gte=time_from)
            except ValueError:
                return Response({
                    'error': 'Invalid time_from format. Use HH:MM format.'
                }, status=status.HTTP_400_BAD_REQUEST)
        
        time_to = request.query_params.get('time_to')
        if time_to:
            try:
                from datetime import time
                time_to = datetime.strptime(time_to, '%H:%M').time()
                queryset = queryset.filter(end_time__lte=time_to)
            except ValueError:
                return Response({
                    'error': 'Invalid time_to format. Use HH:MM format.'
                }, status=status.HTTP_400_BAD_REQUEST)
        
        # Patient and Doctor ID filters
        patient_id = request.query_params.get('patient_id')
        if patient_id:
            queryset = queryset.filter(patient_id=patient_id)
        
        doctor_id = request.query_params.get('doctor_id')
        if doctor_id:
            queryset = queryset.filter(doctor_id=doctor_id)
        
        # Subscription package filter
        subscription_package = request.query_params.get('subscription_package')
        if subscription_package:
            queryset = queryset.filter(
                patient__patientsubscription__package__name__icontains=subscription_package,
                patient__patientsubscription__status='active'
            )
        
        # Boolean filters
        has_diagnosis = request.query_params.get('has_diagnosis')
        if has_diagnosis is not None:
            if has_diagnosis.lower() == 'true':
                queryset = queryset.exclude(diagnosis__isnull=True).exclude(diagnosis='')
            elif has_diagnosis.lower() == 'false':
                queryset = queryset.filter(Q(diagnosis__isnull=True) | Q(diagnosis=''))
        
        has_treatment = request.query_params.get('has_treatment')
        if has_treatment is not None:
            if has_treatment.lower() == 'true':
                queryset = queryset.exclude(treatment__isnull=True).exclude(treatment='')
            elif has_treatment.lower() == 'false':
                queryset = queryset.filter(Q(treatment__isnull=True) | Q(treatment=''))
        
        # Ordering
        ordering = request.query_params.get('ordering', '-appointment_date')
        valid_orderings = [
            'appointment_date', '-appointment_date',
            'start_time', '-start_time',
            'status', '-status',
            'appointment_type', '-appointment_type',
            'created_at', '-created_at',
            'updated_at', '-updated_at'
        ]
        
        if ordering in valid_orderings:
            queryset = queryset.order_by(ordering)
        elif ordering in ['patient_name', '-patient_name']:
            order_direction = '' if ordering == 'patient_name' else '-'
            queryset = queryset.order_by(f'{order_direction}patient__user__first_name', f'{order_direction}patient__user__last_name')
        elif ordering in ['doctor_name', '-doctor_name']:
            order_direction = '' if ordering == 'doctor_name' else '-'
            queryset = queryset.order_by(f'{order_direction}doctor__user__first_name', f'{order_direction}doctor__user__last_name')
        else:
            # Default ordering
            queryset = queryset.order_by('-appointment_date', '-start_time')
        
        # Select related to optimize queries
        queryset = queryset.select_related(
            'patient__user',
            'doctor__user', 
            'healthcare_facility'
        ).prefetch_related(
            'patient__patientsubscription_set__package'
        )
        
        # Paginate results
        page = self.paginate_queryset(queryset)
        if page is not None:
            serializer = self.get_serializer(page, many=True)
            return self.get_paginated_response(serializer.data)
        
        serializer = self.get_serializer(queryset, many=True)
        return Response(serializer.data)


# class AppointmentDocumentViewSet(viewsets.ModelViewSet):
#     queryset = AppointmentDocument.objects.all()
#     serializer_class = AppointmentDocumentSerializer
#     permission_classes = [permissions.IsAuthenticated]
#     
#     def get_permissions(self):
#         if self.action in ['create']:
#             permission_classes = [IsDoctorUser | IsAdminUser]
#         elif self.action in ['destroy']:
#             permission_classes = [IsAdminUser]
#         else:
#             permission_classes = [IsPatientOrDoctorOrAdmin]
#         return [permission() for permission in permission_classes]
#     
#     def get_queryset(self):
#         queryset = AppointmentDocument.objects.all()
#         
#         # Filter by role
#         if self.request.user.is_authenticated and self.request.user.roles.filter(name='admin').exists():
#             # Admin can see all documents
#             pass
#         elif self.request.user.is_authenticated and self.request.user.roles.filter(name='doctor').exists():
#             # Doctors can only see documents for their appointments
#             try:
#                 doctor = self.request.user.doctor
#                 queryset = queryset.filter(appointment__doctor=doctor)
#             except Doctor.DoesNotExist:
#                 return AppointmentDocument.objects.none()
#         elif self.request.user.is_authenticated and self.request.user.roles.filter(name='patient').exists():
#             # Patients can only see their own documents
#             try:
#                 patient = self.request.user.patient
#                 queryset = queryset.filter(appointment__patient=patient)
#             except Patient.DoesNotExist:
#                 return AppointmentDocument.objects.none()
#         else:
#             return AppointmentDocument.objects.none()
#             
#         # Optional filtering
#         appointment_id = self.request.query_params.get('appointment_id')
#         if appointment_id:
#             queryset = queryset.filter(appointment_id=appointment_id)
#             
#         document_type = self.request.query_params.get('type')
#         if document_type:
#             queryset = queryset.filter(document_type=document_type)
#                 
#         return queryset
#     
#     def perform_create(self, serializer):
#         # Set uploaded_by to current user
#         serializer.save(uploaded_by=self.request.user)
# 
# 
from .twilio_utils import create_twilio_room, close_twilio_room, generate_twilio_token
# import uuid
# 
class ConsultationViewSet(viewsets.ModelViewSet):
    queryset = Consultation.objects.all()
    serializer_class = ConsultationSerializer
    permission_classes = [permissions.IsAuthenticated]
    
    def get_permissions(self):
        if self.action in ['create']:
            permission_classes = [IsDoctorUser | IsAdminUser]
        elif self.action in ['destroy']:
            permission_classes = [IsAdminUser]
        else:
            permission_classes = [IsPatientOrDoctorOrAdmin]
        return [permission() for permission in permission_classes]
    
    def get_queryset(self):
        queryset = Consultation.objects.all()
        
        # Filter by role
        if self.request.user.is_authenticated and self.request.user.roles.filter(name='admin').exists():
            # Admin can see all consultations
            pass
        elif self.request.user.is_authenticated and self.request.user.roles.filter(name='doctor').exists():
            # Doctors can only see their own consultations
            try:
                doctor = self.request.user.doctor
                queryset = queryset.filter(appointment__doctor=doctor)
            except Doctor.DoesNotExist:
                return Consultation.objects.none()
        elif self.request.user.is_authenticated and self.request.user.roles.filter(name='patient').exists():
            # Patients can only see their own consultations
            try:
                patient = self.request.user.patient
                queryset = queryset.filter(appointment__patient=patient)
            except Patient.DoesNotExist:
                return Consultation.objects.none()
        else:
            return Consultation.objects.none()
            
        # Optional filtering
        appointment_id = self.request.query_params.get('appointment_id')
        if appointment_id:
            queryset = queryset.filter(appointment_id=appointment_id)
            
        status_param = self.request.query_params.get('status')
        if status_param:
            queryset = queryset.filter(status=status_param)
                
        return queryset

    @swagger_auto_schema(
        operation_description="Start a video consultation with a patient",
        responses={
            200: openapi.Response("Success", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'id': openapi.Schema(type=openapi.TYPE_STRING, format="uuid"),
                    'appointment': openapi.Schema(type=openapi.TYPE_OBJECT),
                    'status': openapi.Schema(type=openapi.TYPE_STRING),
                    'start_time': openapi.Schema(type=openapi.TYPE_STRING, format="datetime"),
                    'end_time': openapi.Schema(type=openapi.TYPE_STRING, format="datetime"),
                    'session_id': openapi.Schema(type=openapi.TYPE_STRING),
                    'twilio_room_name': openapi.Schema(type=openapi.TYPE_STRING),
                    'twilio_room_sid': openapi.Schema(type=openapi.TYPE_STRING),
                    'token': openapi.Schema(type=openapi.TYPE_STRING, description="Twilio token for doctor")
                }
            )),
            400: openapi.Response("Bad Request", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'error': openapi.Schema(type=openapi.TYPE_STRING)
                }
            )),
            403: openapi.Response("Forbidden", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'error': openapi.Schema(type=openapi.TYPE_STRING)
                }
            )),
            500: openapi.Response("Server Error", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'error': openapi.Schema(type=openapi.TYPE_STRING)
                }
            ))
        }
    )
    @action(detail=True, methods=['post'], permission_classes=[IsDoctorUser])
    def start_consultation(self, request, pk=None):
        """
        Endpoint for doctors to start a consultation and create a Twilio room
        """

        print(request.data)

        consultation = self.get_object()
        print(consultation)

        
        # Check if doctor owns this consultation
        if not consultation.appointment.doctor.user == request.user:
            return Response({
                'error': 'You can only start your own consultations'
            }, status=status.HTTP_403_FORBIDDEN)
            
        # Check if consultation can be started
        if consultation.status not in ['scheduled']:
            return Response({
                'error': f'Cannot start a consultation with status: {consultation.status}'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            # Generate a unique room name
            room_name = f"panacare-consultation-{str(consultation.id)}"
            
            # Start the consultation (update status first)
            consultation.status = 'in-progress'
            consultation.start_time = timezone.now()
            consultation.twilio_room_name = room_name
            consultation.save()
            
            # Update appointment status
            appointment = consultation.appointment
            appointment.status = 'arrived'
            appointment.save()
            
            try:
                # Try to create Twilio room and tokens
                room = create_twilio_room(room_name)
                
                consultation.twilio_room_sid = room.sid
                consultation.session_id = room.sid
                
                # Generate tokens for doctor and patient
                doctor_identity = f"doctor-{consultation.appointment.doctor.id}"
                patient_identity = f"patient-{consultation.appointment.patient.id}"

                print("This is the doctor identity: ", doctor_identity)
                print("This is the patient identity: ", patient_identity)
                
                consultation.doctor_token = generate_twilio_token(doctor_identity, room_name)
                consultation.patient_token = generate_twilio_token(patient_identity, room_name)
                consultation.save()
                
                # Return the consultation data with doctor token
                serializer = self.get_serializer(consultation)
                response_data = serializer.data
                response_data['token'] = consultation.doctor_token
                
                return Response(response_data)
                
            except Exception as twilio_error:
                # If Twilio fails, still allow consultation to proceed without video
                consultation.save()
                
                serializer = self.get_serializer(consultation)
                response_data = serializer.data
                response_data['warning'] = f'Consultation started but video calling unavailable: {str(twilio_error)}'
                response_data['token'] = None
                
                return Response(response_data)
                
        except Exception as e:
            # Check if it's a Twilio credentials error
            error_msg = str(e)
            if "Authentication Error" in error_msg or "credentials" in error_msg.lower():
                return Response({
                    'error': f'Failed to start consultation: Twilio credentials not configured properly - {str(e)}'
                }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            else:
                return Response({
                    'error': f'Failed to start consultation: {str(e)}'
                }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
    @action(detail=True, methods=['post'], permission_classes=[IsDoctorUser])
    def end_consultation(self, request, pk=None):
        """
        Endpoint for doctors to end a consultation and close the Twilio room
        """
        consultation = self.get_object()
        
        # Check if doctor owns this consultation
        if not consultation.appointment.doctor.user == request.user:
            return Response({
                'error': 'You can only end your own consultations'
            }, status=status.HTTP_403_FORBIDDEN)
            
        # Check if consultation can be ended
        if consultation.status not in ['in-progress']:
            return Response({
                'error': f'Cannot end a consultation with status: {consultation.status}'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            # If there's a Twilio room, close it
            if consultation.twilio_room_sid:
                close_twilio_room(consultation.twilio_room_sid)
            
            # End the consultation
            consultation.status = 'completed'
            consultation.end_time = timezone.now()
            consultation.save()
            
            # Update appointment status
            appointment = consultation.appointment
            appointment.status = 'fulfilled'
            appointment.save()
            
            serializer = self.get_serializer(consultation)
            return Response(serializer.data)
            
        except Exception as e:
            return Response({
                'error': f'Failed to end consultation: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
    @action(detail=True, methods=['get'], permission_classes=[IsPatientOrDoctorOrAdmin])
    def get_token(self, request, pk=None):
        """
        Endpoint for patients and doctors to get their Twilio token
        """
        consultation = self.get_object()
        
        # Check if the user is either the patient or doctor for this consultation
        is_doctor = request.user.roles.filter(name='doctor').exists() and hasattr(request.user, 'doctor') and request.user.doctor == consultation.appointment.doctor
        is_patient = request.user.roles.filter(name='patient').exists() and hasattr(request.user, 'patient') and request.user.patient == consultation.appointment.patient
        is_admin = request.user.roles.filter(name='admin').exists()
        
        if not (is_doctor or is_patient or is_admin):
            return Response({
                'error': 'You are not authorized to access this consultation'
            }, status=status.HTTP_403_FORBIDDEN)
        
        # Check if consultation is in progress
        if consultation.status != 'in-progress':
            return Response({
                'error': f'Cannot join a consultation with status: {consultation.status}'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            token = None
            now = timezone.now()
            expired = now >= consultation.created_at + timedelta(hours=1)

            # Doctor logic
            if is_doctor:
                token = consultation.doctor_token
                if not token or expired:
                    identity = f"doctor-{consultation.appointment.doctor.id}"
                    room_name = consultation.twilio_room_name
                    try:
                        token = generate_twilio_token(identity, room_name)
                        consultation.doctor_token = token
                        consultation.created_at = now
                        consultation.save(update_fields=['doctor_token', 'created_at'])
                    except Exception as twilio_error:
                        return Response({
                            'error': f'Video calling unavailable: {str(twilio_error)}',
                            'consultation_active': True,
                            'room_name': consultation.twilio_room_name
                        }, status=status.HTTP_503_SERVICE_UNAVAILABLE)

            # Patient logic
            elif is_patient:
                token = consultation.patient_token
                if not token or expired:
                    identity = f"patient-{consultation.appointment.patient.id}"
                    room_name = consultation.twilio_room_name
                    try:
                        token = generate_twilio_token(identity, room_name)
                        consultation.patient_token = token
                        consultation.created_at = now
                        consultation.save(update_fields=['patient_token', 'created_at'])
                    except Exception as twilio_error:
                        return Response({
                            'error': f'Video calling unavailable: {str(twilio_error)}',
                            'consultation_active': True,
                            'room_name': consultation.twilio_room_name
                        }, status=status.HTTP_503_SERVICE_UNAVAILABLE)

            return Response({
                'token': token,
                'room_name': consultation.twilio_room_name
            })

        except Exception as e:
            return Response({
                'error': f'Failed to get token: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            
    @swagger_auto_schema(
        operation_description="Join or rejoin an ongoing consultation as a patient or doctor",
        responses={
            200: openapi.Response("Success", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'token': openapi.Schema(type=openapi.TYPE_STRING, description="Twilio token"),
                    'room_name': openapi.Schema(type=openapi.TYPE_STRING, description="Twilio room name")
                }
            )),
            400: openapi.Response("Bad Request", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'error': openapi.Schema(type=openapi.TYPE_STRING)
                }
            )),
            403: openapi.Response("Forbidden", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'error': openapi.Schema(type=openapi.TYPE_STRING)
                }
            )),
            500: openapi.Response("Server Error", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'error': openapi.Schema(type=openapi.TYPE_STRING)
                }
            ))
        }
    )
    @action(detail=True, methods=['post'], permission_classes=[IsPatientOrDoctorOrAdmin])
    def join_consultation(self, request, pk=None):
        """
        Endpoint for patients or doctors to join/rejoin a consultation
        """
        consultation = self.get_object()
        
        # Check if the user is the patient or doctor for this consultation
        is_patient = request.user.roles.filter(name='patient').exists() and hasattr(request.user, 'patient') and request.user.patient == consultation.appointment.patient
        is_doctor = request.user.roles.filter(name='doctor').exists() and hasattr(request.user, 'doctor') and request.user.doctor == consultation.appointment.doctor
        
        if not (is_patient or is_doctor):
            return Response({
                'error': 'You can only join consultations you are a participant in'
            }, status=status.HTTP_403_FORBIDDEN)
        
        # Check if consultation is in progress
        if consultation.status != 'in-progress':
            return Response({
                'error': f'Cannot join a consultation with status: {consultation.status}'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            # Generate or retrieve token based on user role
            token = None
            identity = None
            
            if is_doctor:
                token = consultation.doctor_token
                identity = f"doctor-{consultation.appointment.doctor.id}"
            elif is_patient:
                token = consultation.patient_token
                identity = f"patient-{consultation.appointment.patient.id}"
            
            # If token doesn't exist or is expired, generate a new one
            if not token:
                room_name = consultation.twilio_room_name
                token = generate_twilio_token(identity, room_name)
                
                # Save the new token
                if is_doctor:
                    consultation.doctor_token = token
                elif is_patient:
                    consultation.patient_token = token
                    
                consultation.save()
            
            return Response({
                'token': token,
                'room_name': consultation.twilio_room_name
            })
            
        except Exception as e:
            return Response({
                'error': f'Failed to join consultation: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            
    @action(detail=True, methods=['get'], permission_classes=[IsPatientOrDoctorOrAdmin])
    def chat_messages(self, request, pk=None):
        """
        Endpoint to get chat messages for a consultation
        """
        consultation = self.get_object()
        
        # Check if the user is a participant in this consultation
        is_patient = request.user.roles.filter(name='patient').exists() and hasattr(request.user, 'patient') and request.user.patient == consultation.appointment.patient
        is_doctor = request.user.roles.filter(name='doctor').exists() and hasattr(request.user, 'doctor') and request.user.doctor == consultation.appointment.doctor
        is_admin = request.user.roles.filter(name='admin').exists()
        
        if not (is_patient or is_doctor or is_admin):
            return Response({
                'error': 'You can only access chat messages for consultations you are a participant in'
            }, status=status.HTTP_403_FORBIDDEN)
        
        # Get messages with optional limit
        limit = request.query_params.get('limit', 50)
        try:
            limit = int(limit)
        except ValueError:
            limit = 50
            
        # Get messages
        messages = ConsultationChat.objects.filter(consultation=consultation).order_by('-created_at')[:limit]
        serializer = ConsultationChatSerializer(messages, many=True)
        
        return Response(serializer.data)
    
    @swagger_auto_schema(
        operation_description="Send a chat message in a consultation",
        request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            required=['message'],
            properties={
                'message': openapi.Schema(type=openapi.TYPE_STRING, description="Chat message content")
            }
        ),
        responses={
            200: ConsultationChatSerializer,
            400: openapi.Response("Bad Request", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'error': openapi.Schema(type=openapi.TYPE_STRING)
                }
            )),
            403: openapi.Response("Forbidden", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'error': openapi.Schema(type=openapi.TYPE_STRING)
                }
            ))
        }
    )
    @action(detail=True, methods=['post'], permission_classes=[IsPatientOrDoctorOrAdmin])
    def send_message(self, request, pk=None):
        """
        Endpoint to send a chat message in a consultation
        """
        consultation = self.get_object()
        
        # Check if the user is a participant in this consultation
        is_patient = request.user.roles.filter(name='patient').exists() and hasattr(request.user, 'patient') and request.user.patient == consultation.appointment.patient
        is_doctor = request.user.roles.filter(name='doctor').exists() and hasattr(request.user, 'doctor') and request.user.doctor == consultation.appointment.doctor
        
        if not (is_patient or is_doctor):
            return Response({
                'error': 'You can only send messages in consultations you are a participant in'
            }, status=status.HTTP_403_FORBIDDEN)
        
        # Check if consultation is in progress
        if consultation.status != 'in-progress':
            return Response({
                'error': f'Cannot send messages in a consultation with status: {consultation.status}'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # Get message content
        message_content = request.data.get('message')
        if not message_content:
            return Response({
                'error': 'Message content is required'
            }, status=status.HTTP_400_BAD_REQUEST)
            
        # Create the message
        is_doctor = request.user.roles.filter(name='doctor').exists()
        
        chat_message = ConsultationChat.objects.create(
            consultation=consultation,
            message=message_content,
            sender=request.user,
            is_doctor=is_doctor
        )
        
        serializer = ConsultationChatSerializer(chat_message)
        return Response(serializer.data)
    
    @action(detail=True, methods=['post'], permission_classes=[IsPatientOrDoctorOrAdmin])
    def mark_messages_read(self, request, pk=None):
        """
        Endpoint to mark all unread messages as read
        """
        consultation = self.get_object()
        
        # Check if the user is a participant in this consultation
        is_patient = request.user.roles.filter(name='patient').exists() and hasattr(request.user, 'patient') and request.user.patient == consultation.appointment.patient
        is_doctor = request.user.roles.filter(name='doctor').exists() and hasattr(request.user, 'doctor') and request.user.doctor == consultation.appointment.doctor
        
        if not (is_patient or is_doctor):
            return Response({
                'error': 'You can only mark messages as read in consultations you are a participant in'
            }, status=status.HTTP_403_FORBIDDEN)
            
        # Get all unread messages sent by the other participant
        if is_doctor:
            unread_messages = ConsultationChat.objects.filter(
                consultation=consultation,
                is_read=False,
                is_doctor=False  # Messages from patient
            )
        else:
            unread_messages = ConsultationChat.objects.filter(
                consultation=consultation,
                is_read=False,
                is_doctor=True  # Messages from doctor
            )
            
        # Mark them as read
        now = timezone.now()
        count = unread_messages.count()
        unread_messages.update(is_read=True, read_at=now)
        
        return Response({
            'status': 'success',
            'messages_marked_read': count
        })


# class PackageViewSet(viewsets.ModelViewSet):
#     queryset = Package.objects.all()
#     serializer_class = PackageSerializer
#     permission_classes = [permissions.IsAuthenticated]
#     
#     def get_permissions(self):
#         if self.action in ['create', 'update', 'partial_update', 'destroy']:
#             permission_classes = [IsAdminUser]
#         else:
#             permission_classes = [permissions.IsAuthenticated]
#         return [permission() for permission in permission_classes]
#     
#     def get_queryset(self):
#         queryset = Package.objects.filter(is_active=True)
#         
#         # Admin can see inactive packages too with a filter
#         if self.request.user.is_authenticated and self.request.user.roles.filter(name='admin').exists():
#             include_inactive = self.request.query_params.get('include_inactive')
#             if include_inactive and include_inactive.lower() == 'true':
#                 queryset = Package.objects.all()
#                 
#         return queryset
# 
# 
# class PatientSubscriptionViewSet(viewsets.ModelViewSet):
#     queryset = PatientSubscription.objects.all()
#     serializer_class = PatientSubscriptionSerializer
#     permission_classes = [permissions.IsAuthenticated]
#     
#     def get_permissions(self):
#         if self.action in ['create']:
#             permission_classes = [IsPatientUser | IsAdminUser]
#         elif self.action in ['update', 'partial_update', 'destroy']:
#             permission_classes = [IsAdminUser]
#         else:
#             permission_classes = [permissions.IsAuthenticated]
#         return [permission() for permission in permission_classes]
#     
#     def get_queryset(self):
#         queryset = PatientSubscription.objects.all()
#         
#         # Filter by role
#         if self.request.user.is_authenticated and self.request.user.roles.filter(name='admin').exists():
#             # Admin can see all subscriptions
#             pass
#         elif self.request.user.is_authenticated and self.request.user.roles.filter(name='patient').exists():
#             # Patients can only see their own subscriptions
#             try:
#                 patient = self.request.user.patient
#                 queryset = queryset.filter(patient=patient)
#             except Patient.DoesNotExist:
#                 return PatientSubscription.objects.none()
#         else:
#             return PatientSubscription.objects.none()
#             
#         # Optional filtering
#         patient_id = self.request.query_params.get('patient_id')
#         if patient_id:
#             queryset = queryset.filter(patient_id=patient_id)
#             
#         package_id = self.request.query_params.get('package_id')
#         if package_id:
#             queryset = queryset.filter(package_id=package_id)
#             
#         status_param = self.request.query_params.get('status')
#         if status_param:
#             queryset = queryset.filter(status=status_param)
#                 
#         return queryset
#     
#     @action(detail=False, methods=['get'], permission_classes=[IsPatientUser])
#     def my_subscriptions(self, request):
#         """
#         Endpoint for patients to view their own subscriptions
#         """
#         try:
#             patient = request.user.patient
#             subscriptions = PatientSubscription.objects.filter(patient=patient)
#             
#             # Optional filtering
#             status_param = request.query_params.get('status')
#             if status_param:
#                 subscriptions = subscriptions.filter(status=status_param)
#                 
#             serializer = self.get_serializer(subscriptions, many=True)
#             return Response(serializer.data)
#         except Patient.DoesNotExist:
#             return Response({
#                 'error': 'Patient profile not found'
#             }, status=status.HTTP_404_NOT_FOUND)
#     
#     @action(detail=True, methods=['post'], permission_classes=[IsPatientUser])
#     def cancel_subscription(self, request, pk=None):
#         """
#         Endpoint for patients to cancel their subscription
#         """
#         subscription = self.get_object()
#         
#         # Check if patient owns this subscription
#         if not subscription.patient.user == request.user:
#             return Response({
#                 'error': 'You can only cancel your own subscriptions'
#             }, status=status.HTTP_403_FORBIDDEN)
#             
#         # Check if subscription can be cancelled
#         if subscription.status not in ['active']:
#             return Response({
#                 'error': f'Cannot cancel a subscription with status: {subscription.status}'
#             }, status=status.HTTP_400_BAD_REQUEST)
#             
#         # Cancel the subscription
#         subscription.status = 'cancelled'
#         subscription.save()
#         
#         serializer = self.get_serializer(subscription)
#         return Response(serializer.data)
# 
# 
# class ResourceViewSet(viewsets.ModelViewSet):
#     queryset = Resource.objects.all()
#     serializer_class = ResourceSerializer
#     permission_classes = [permissions.IsAuthenticated]
#     
#     def get_permissions(self):
#         if self.action in ['create']:
#             permission_classes = [IsDoctorUser | IsAdminUser]
#         elif self.action in ['update', 'partial_update', 'destroy']:
#             permission_classes = [IsAdminUser]
#         else:
#             permission_classes = [permissions.IsAuthenticated]
#         return [permission() for permission in permission_classes]
#     
#     def get_queryset(self):
#         queryset = Resource.objects.filter(is_active=True, is_approved=True)
#         
#         # Admin can see all resources
#         if self.request.user.is_authenticated and self.request.user.roles.filter(name='admin').exists():
#             include_unapproved = self.request.query_params.get('include_unapproved')
#             if include_unapproved and include_unapproved.lower() == 'true':
#                 queryset = Resource.objects.all()
#         
#         # Doctors can see their own unapproved resources
#         elif self.request.user.is_authenticated and self.request.user.roles.filter(name='doctor').exists():
#             try:
#                 doctor = self.request.user.doctor
#                 include_unapproved = self.request.query_params.get('include_unapproved')
#                 if include_unapproved and include_unapproved.lower() == 'true':
#                     queryset = queryset.filter(Q(is_approved=True) | Q(author=doctor))
#             except Doctor.DoesNotExist:
#                 pass
#                 
#         # Optional filtering
#         category = self.request.query_params.get('category')
#         if category:
#             queryset = queryset.filter(category=category)
#             
#         content_type = self.request.query_params.get('content_type')
#         if content_type:
#             queryset = queryset.filter(content_type=content_type)
#             
#         search = self.request.query_params.get('search')
#         if search:
#             queryset = queryset.filter(
#                 Q(title__icontains=search) | 
#                 Q(description__icontains=search) |
#                 Q(tags__icontains=search)
#             )
#                 
#         return queryset
#     
#     def perform_create(self, serializer):
#         # Set author to doctor if user is a doctor
#         if self.request.user.roles.filter(name='doctor').exists():
#             try:
#                 doctor = self.request.user.doctor
#                 serializer.save(author=doctor, is_approved=False)
#             except Doctor.DoesNotExist:
#                 serializer.save(is_approved=False)
#         else:
#             # Admin uploads are auto-approved
#             serializer.save(is_approved=True, approved_by=self.request.user)
#     
#     @action(detail=True, methods=['post'], permission_classes=[IsAdminUser])
#     def approve_resource(self, request, pk=None):
#         """
#         Endpoint for admins to approve a resource
#         """
#         resource = self.get_object()
#         
#         # Check if resource is already approved
#         if resource.is_approved:
#             return Response({
#                 'error': 'Resource is already approved'
#             }, status=status.HTTP_400_BAD_REQUEST)
#             
#         # Approve the resource
#         resource.is_approved = True
#         resource.approved_by = request.user
#         resource.save()
#         
#         serializer = self.get_serializer(resource)
#         return Response(serializer.data)
#     
#     @action(detail=True, methods=['post'], permission_classes=[permissions.IsAuthenticated])
#     def verify_password(self, request, pk=None):
#         """
#         Endpoint to verify password for password-protected resources
#         """
#         resource = self.get_object()
#         
#         # Check if resource is password-protected
#         if not resource.is_password_protected:
#             return Response({
#                 'error': 'This resource is not password-protected'
#             }, status=status.HTTP_400_BAD_REQUEST)
#             
#         # Verify password
#         from django.contrib.auth.hashers import check_password
#         password = request.data.get('password')
#         if password and check_password(password, resource.password_hash):
#             return Response({
#                 'message': 'Password verified',
#                 'resource': self.get_serializer(resource).data
#             })
#         else:
#             return Response({
#                 'error': 'Invalid password'
#             }, status=status.HTTP_401_UNAUTHORIZED)
# 
# 
class DoctorRatingViewSet(viewsets.ModelViewSet):
    queryset = DoctorRating.objects.all()
    serializer_class = DoctorRatingSerializer
    permission_classes = [permissions.IsAuthenticated]
    
    def get_permissions(self):
        if self.action in ['create']:
            permission_classes = [IsPatientUser]
        elif self.action in ['destroy']:
            permission_classes = [IsAdminUser]
        else:
            permission_classes = [permissions.IsAuthenticated]
        return [permission() for permission in permission_classes]
    
    def get_queryset(self):
        queryset = DoctorRating.objects.all()
        
        # Optional filtering
        doctor_id = self.request.query_params.get('doctor_id')
        if doctor_id:
            queryset = queryset.filter(doctor_id=doctor_id)
            
        patient_id = self.request.query_params.get('patient_id')
        if patient_id:
            queryset = queryset.filter(patient_id=patient_id)
            
        rating = self.request.query_params.get('rating')
        if rating:
            try:
                rating = int(rating)
                queryset = queryset.filter(rating=rating)
            except ValueError:
                pass
                
        return queryset
    
    def perform_create(self, serializer):
        try:
            # Set patient to current user's patient
            patient = self.request.user.patient
            
            # Check if patient has already rated this doctor
            doctor_id = serializer.validated_data.get('doctor').id
            existing_rating = DoctorRating.objects.filter(doctor_id=doctor_id, patient=patient).exists()
            
            if existing_rating:
                # Update the existing rating
                existing_rating = DoctorRating.objects.get(doctor_id=doctor_id, patient=patient)
                existing_rating.rating = serializer.validated_data.get('rating')
                existing_rating.review = serializer.validated_data.get('review', '')
                existing_rating.is_anonymous = serializer.validated_data.get('is_anonymous', False)
                existing_rating.save()
            else:
                # Create new rating
                serializer.save(patient=patient)
        except Patient.DoesNotExist:
            raise serializers.ValidationError("Patient profile not found")
    
    @action(detail=False, methods=['get'])
    def doctor_average_rating(self, request):
        """
        Get the average rating for a doctor
        """
        doctor_id = request.query_params.get('doctor_id')
        if not doctor_id:
            return Response({
                'error': 'doctor_id parameter is required'
            }, status=status.HTTP_400_BAD_REQUEST)
            
        # Get average rating
        average_rating = DoctorRating.objects.filter(doctor_id=doctor_id).aggregate(Avg('rating'))
        total_ratings = DoctorRating.objects.filter(doctor_id=doctor_id).count()
        
        return Response({
            'doctor_id': doctor_id,
            'average_rating': average_rating['rating__avg'] or 0,
            'total_ratings': total_ratings
        })


class ArticleViewSet(viewsets.ModelViewSet):
    queryset = Article.objects.all()
    serializer_class = ArticleSerializer
    permission_classes = [permissions.IsAuthenticated]
    
    @swagger_auto_schema(
        operation_description="""
 List all articles with filtering options.
 
 This endpoint fetches articles based on user permissions:
 - Admins: Can see all articles, including unapproved and unpublished
 - Doctors: Can see all approved/published articles plus their own drafts
 - Patients with active subscription: Can see public and subscriber-only articles
 - Other users: Can see only public articles
 
 Examples:
 - Get all articles: GET /api/articles/
 - Get articles in nutrition category: GET /api/articles/?category=nutrition
 - Get featured articles: GET /api/articles/?featured=true
 - Get articles by condition: GET /api/articles/?condition=diabetes
 - Get articles sorted by popularity: GET /api/articles/?sort_by=popular
 - Search articles: GET /api/articles/?search=diabetes
        """,
        manual_parameters=[
            openapi.Parameter('category', openapi.IN_QUERY, type=openapi.TYPE_STRING, description="Filter by category (general, nutrition, fitness, mental, children, chronic, prevention, research, other)"),
            openapi.Parameter('author_id', openapi.IN_QUERY, type=openapi.TYPE_STRING, description="Filter by author ID"),
            openapi.Parameter('featured', openapi.IN_QUERY, type=openapi.TYPE_BOOLEAN, description="Filter by featured status"),
            openapi.Parameter('visibility', openapi.IN_QUERY, type=openapi.TYPE_STRING, enum=["public", "subscribers", "private"], description="Filter by visibility level"),
            openapi.Parameter('condition', openapi.IN_QUERY, type=openapi.TYPE_STRING, description="Filter by related health condition"),
            openapi.Parameter('date_from', openapi.IN_QUERY, type=openapi.TYPE_STRING, format="date", description="Filter by publish date (from)"),
            openapi.Parameter('date_to', openapi.IN_QUERY, type=openapi.TYPE_STRING, format="date", description="Filter by publish date (to)"),
            openapi.Parameter('sort_by', openapi.IN_QUERY, type=openapi.TYPE_STRING, enum=["popular", "newest", "oldest"], description="Sort by view count (popular), newest publish date (newest), or oldest publish date (oldest)"),
            openapi.Parameter('search', openapi.IN_QUERY, type=openapi.TYPE_STRING, description="Search across title, content, summary, tags, and related_conditions"),
            openapi.Parameter('is_approved', openapi.IN_QUERY, type=openapi.TYPE_BOOLEAN, description="Filter by approval status (admin only)"),
            openapi.Parameter('is_published', openapi.IN_QUERY, type=openapi.TYPE_BOOLEAN, description="Filter by publication status"),
            openapi.Parameter('reading_time', openapi.IN_QUERY, type=openapi.TYPE_INTEGER, description="Filter by reading time in minutes (returns articles with reading time <= specified value)"),
            openapi.Parameter('limit', openapi.IN_QUERY, type=openapi.TYPE_INTEGER, description="Limit number of results returned")
        ],
        responses={
            200: ArticleSerializer(many=True)
        }
    )
    def list(self, request, *args, **kwargs):
        return super().list(request, *args, **kwargs)
    
    def retrieve(self, request, *args, **kwargs):
        """
        Custom retrieve method to handle article access permissions
        """
        instance = self.get_object()
        
        # Check if user has permission to view this article
        if self.request.user.is_authenticated and self.request.user.roles.filter(name='admin').exists():
            # Admin can view any article
            pass
        elif self.request.user.is_authenticated and self.request.user.roles.filter(name='doctor').exists():
            # Doctors can view all published articles or their own articles
            try:
                doctor = self.request.user.doctor
                if not (instance.is_approved and instance.is_published) and instance.author != doctor:
                    # Article is not published and doctor is not the author
                    return Response({
                        'error': 'Article not found or access denied'
                    }, status=status.HTTP_404_NOT_FOUND)
            except Doctor.DoesNotExist:
                # If doctor profile doesn't exist, only allow published articles
                if not (instance.is_approved and instance.is_published):
                    return Response({
                        'error': 'Article not found or access denied'
                    }, status=status.HTTP_404_NOT_FOUND)
        else:
            # Patients and other users
            if not (instance.is_approved and instance.is_published):
                return Response({
                    'error': 'Article not found or access denied'
                }, status=status.HTTP_404_NOT_FOUND)
            
            # Check visibility permissions for patients
            if instance.visibility == 'private':
                return Response({
                    'error': 'Article not found or access denied'
                }, status=status.HTTP_404_NOT_FOUND)
            elif instance.visibility == 'subscribers':
                # Check if patient has active subscription
                has_active_subscription = False
                if hasattr(self.request.user, 'patient'):
                    try:
                        patient = self.request.user.patient
                        has_active_subscription = PatientSubscription.objects.filter(
                            patient=patient,
                            status='active',
                            end_date__gte=timezone.now().date()
                        ).exists()
                    except Exception:
                        has_active_subscription = False
                
                if not has_active_subscription:
                    return Response({
                        'error': 'This article is only available to subscribers'
                    }, status=status.HTTP_403_FORBIDDEN)
        
        # Update view count
        instance.view_count += 1
        instance.save(update_fields=['view_count'])
        
        serializer = self.get_serializer(instance)
        return Response(serializer.data)
    
    @swagger_auto_schema(
        operation_description="Create a new article (doctors only). Note: featured_image should be sent as a file upload in a multipart form request.",
        request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            required=['title', 'content', 'category'],
            properties={
                'title': openapi.Schema(type=openapi.TYPE_STRING, description='Article title'),
                'content': openapi.Schema(type=openapi.TYPE_STRING, description='Article content'),
                'summary': openapi.Schema(type=openapi.TYPE_STRING, description='Brief summary of the article'),
                'category': openapi.Schema(type=openapi.TYPE_STRING, description='Article category'),
                'tags': openapi.Schema(type=openapi.TYPE_STRING, description='Comma-separated tags'),
                'featured_image': openapi.Schema(type=openapi.TYPE_FILE, description='Featured image for the article'),
                'visibility': openapi.Schema(type=openapi.TYPE_STRING, enum=['public', 'subscribers', 'private'], description='Controls who can view this article'),
                'related_conditions': openapi.Schema(type=openapi.TYPE_STRING, description='Comma-separated health conditions'),
                'reading_time': openapi.Schema(type=openapi.TYPE_INTEGER, description='Estimated reading time in minutes')
            }
        ),
        responses={
            201: ArticleSerializer,
            400: openapi.Response("Bad Request", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'error': openapi.Schema(type=openapi.TYPE_STRING)
                }
            ))
        }
    )
    def create(self, request, *args, **kwargs):
        return super().create(request, *args, **kwargs)
    
    def get_permissions(self):
        if self.action in ['create']:
            permission_classes = [IsDoctorUser]
        elif self.action in ['update', 'partial_update']:
            permission_classes = [IsDoctorUser]  # Only author can edit, enforced in perform_update
        elif self.action in ['destroy']:
            permission_classes = [IsDoctorUser]  # Allow doctors to delete their own articles
        else:
            permission_classes = [permissions.IsAuthenticated]
        return [permission() for permission in permission_classes]
    
    def get_queryset(self):
        # Basic queryset filtering
        if self.request.user.is_authenticated and self.request.user.roles.filter(name='admin').exists():
            # Admins can see all articles including unapproved ones
            queryset = Article.objects.all()
        elif self.request.user.is_authenticated and self.request.user.roles.filter(name='doctor').exists():
            # Doctors can see all approved/published articles plus their own drafts
            try:
                doctor = self.request.user.doctor
                queryset = Article.objects.filter(
                    Q(is_approved=True, is_published=True) |  # All published articles
                    Q(author=doctor)  # Their own articles regardless of approval/publish status
                )
            except Doctor.DoesNotExist:
                # If doctor profile doesn't exist, fall back to published articles only
                queryset = Article.objects.filter(is_approved=True, is_published=True)
        else:
            # Patients and other users can only see approved and published articles
            queryset = Article.objects.filter(is_approved=True, is_published=True)
            
            # Apply visibility filtering for patients
            # If patient has an active subscription, they can see 'subscribers' articles
            # Otherwise, they can only see 'public' articles
            has_active_subscription = False
            has_patient_profile = hasattr(self.request.user, 'patient')
            
            if has_patient_profile:
                # Case 1: User has a patient profile, check for active subscription
                try:
                    patient = self.request.user.patient
                    has_active_subscription = PatientSubscription.objects.filter(
                        patient=patient,
                        status='active',
                        end_date__gte=timezone.now().date()
                    ).exists()
                except Exception:
                    # Case 2: Error getting subscription info, default to no subscription
                    has_active_subscription = False
            
            # Case 3: Normal subscriber with active subscription can see subscriber content
            # Users without patient profile or without active subscription see only public articles
            if has_active_subscription:
                # Subscribers can see both public and subscriber-only content
                queryset = queryset.filter(visibility__in=['public', 'subscribers'])
            else:
                # Filter to only show public articles to non-subscribers
                queryset = queryset.filter(visibility='public')
            
        # Additional filtering
        category = self.request.query_params.get('category')
        if category:
            queryset = queryset.filter(category=category)
            
        author_id = self.request.query_params.get('author_id')
        if author_id:
            queryset = queryset.filter(author_id=author_id)
            
        # Filter by featured status
        featured = self.request.query_params.get('featured')
        if featured and featured.lower() == 'true':
            queryset = queryset.filter(is_featured=True)
            
        # Filter by visibility
        visibility = self.request.query_params.get('visibility')
        if visibility:
            queryset = queryset.filter(visibility=visibility)
            
        # Filter by related health condition
        condition = self.request.query_params.get('condition')
        if condition:
            queryset = queryset.filter(related_conditions__icontains=condition)
            
        # Date range filtering
        date_from = self.request.query_params.get('date_from')
        if date_from:
            try:
                date_from = datetime.strptime(date_from, '%Y-%m-%d').date()
                queryset = queryset.filter(publish_date__date__gte=date_from)
            except ValueError:
                pass
                
        date_to = self.request.query_params.get('date_to')
        if date_to:
            try:
                date_to = datetime.strptime(date_to, '%Y-%m-%d').date()
                queryset = queryset.filter(publish_date__date__lte=date_to)
            except ValueError:
                pass
            
        # Sort by options
        sort_by = self.request.query_params.get('sort_by')
        if sort_by:
            if sort_by == 'popular':
                queryset = queryset.order_by('-view_count')
            elif sort_by == 'newest':
                queryset = queryset.order_by('-publish_date')
            elif sort_by == 'oldest':
                queryset = queryset.order_by('publish_date')
                
        # Free text search across multiple fields
        search = self.request.query_params.get('search')
        if search:
            queryset = queryset.filter(
                Q(title__icontains=search) | 
                Q(content__icontains=search) |
                Q(summary__icontains=search) |
                Q(tags__icontains=search) |
                Q(related_conditions__icontains=search)
            )
            
        # Filter by approval status (admin only)
        if self.request.user.is_authenticated and self.request.user.roles.filter(name='admin').exists():
            is_approved = self.request.query_params.get('is_approved')
            if is_approved:
                queryset = queryset.filter(is_approved=is_approved.lower() == 'true')
                
        # Filter by publication status
        is_published = self.request.query_params.get('is_published')
        if is_published:
            queryset = queryset.filter(is_published=is_published.lower() == 'true')
            
        # Reading time filter
        reading_time = self.request.query_params.get('reading_time')
        if reading_time:
            try:
                reading_time = int(reading_time)
                queryset = queryset.filter(reading_time__lte=reading_time)
            except ValueError:
                pass
            
        return queryset
    
    def get_object(self):
        """
        Override get_object to handle admin/author access for specific actions
        """
        # For admin actions like approve, publish, retrieve, update, delete - 
        # allow admins and authors to access any article (bypass filtering)
        if self.action in ['retrieve', 'update', 'partial_update', 'destroy', 'approve', 'publish', 'unpublish', 'view']:
            # Check if user is admin
            if self.request.user.is_authenticated and self.request.user.roles.filter(name='admin').exists():
                # Admins can access any article
                return get_object_or_404(Article, pk=self.kwargs.get('pk'))
            
            # Check if user is a doctor and this is their own article
            try:
                doctor = self.request.user.doctor
                article = get_object_or_404(Article, pk=self.kwargs.get('pk'))
                
                # For retrieve action, use normal filtering
                if self.action == 'retrieve':
                    return super().get_object()
                
                # For other actions, doctors can only access their own articles
                if article.author == doctor:
                    return article
                else:
                    # For non-authors, fall back to normal filtering
                    return super().get_object()
                    
            except Doctor.DoesNotExist:
                # User is not a doctor, use normal filtering
                return super().get_object()
        
        # For all other actions (list, etc.), use the normal queryset filtering
        return super().get_object()
    
    def perform_create(self, serializer):
        try:
            # Set author to current user's doctor profile
            doctor = self.request.user.doctor
            serializer.save(author=doctor)
        except Doctor.DoesNotExist:
            raise serializers.ValidationError("Doctor profile not found")
    
    def perform_update(self, serializer):
        # Get the article instance
        instance = self.get_object()
        
        # Check if user is the author
        try:
            doctor = self.request.user.doctor
            if instance.author != doctor and not (self.request.user.is_authenticated and self.request.user.roles.filter(name='admin').exists()):
                raise serializers.ValidationError("You can only edit your own articles")
                
            # If this is an already approved article being edited by its author, 
            # reset the approval status
            if instance.is_approved and instance.author == doctor:
                serializer.save(is_approved=False, approval_date=None, approved_by=None)
            else:
                serializer.save()
                
        except Doctor.DoesNotExist:
            # If user is admin, they can edit regardless
            if self.request.user.is_authenticated and self.request.user.roles.filter(name='admin').exists():
                serializer.save()
            else:
                raise serializers.ValidationError("Doctor profile not found")
    
    def perform_destroy(self, instance):
        # Check if user is the author or admin
        try:
            doctor = self.request.user.doctor
            if instance.author != doctor and not (self.request.user.is_authenticated and self.request.user.roles.filter(name='admin').exists()):
                raise serializers.ValidationError("You can only delete your own articles")
            instance.delete()
        except Doctor.DoesNotExist:
            # If user is admin, they can delete regardless
            if self.request.user.is_authenticated and self.request.user.roles.filter(name='admin').exists():
                instance.delete()
            else:
                raise serializers.ValidationError("Doctor profile not found")
    
    @swagger_auto_schema(
        method='post',
        operation_description="Endpoint for admins to approve an article. Can also set publication status, visibility and featured status.",
        request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            properties={
                'approval_notes': openapi.Schema(type=openapi.TYPE_STRING, description='Notes regarding the approval decision'),
                'publish': openapi.Schema(type=openapi.TYPE_STRING, enum=['true', 'false'], description='Whether to also publish the article'),
                'visibility': openapi.Schema(type=openapi.TYPE_STRING, enum=['public', 'subscribers', 'private'], description='Visibility level for the article'),
                'featured': openapi.Schema(type=openapi.TYPE_STRING, enum=['true', 'false'], description='Whether to feature the article')
            }
        ),
        responses={
            200: ArticleSerializer,
            400: openapi.Response("Bad Request", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'error': openapi.Schema(type=openapi.TYPE_STRING)
                }
            ))
        }
    )
    @action(detail=True, methods=['post'], permission_classes=[IsAdminUser])
    def approve(self, request, pk=None):
        """
        Endpoint for admins to approve an article
        """
        article = self.get_object()
        
        # Check if article is already approved
        if article.is_approved:
            return Response({
                'error': 'Article is already approved'
            }, status=status.HTTP_400_BAD_REQUEST)
            
        # Approve the article
        article.is_approved = True
        article.approved_by = request.user
        article.approval_date = timezone.now()
        
        # Add approval notes if provided
        approval_notes = request.data.get('approval_notes')
        if approval_notes:
            article.approval_notes = approval_notes
            
        # Set publish status if provided
        publish = request.data.get('publish')
        if publish and publish.lower() == 'true':
            article.is_published = True
            article.publish_date = timezone.now()
            
        # Set visibility if provided
        visibility = request.data.get('visibility')
        if visibility in ['public', 'subscribers', 'private']:
            article.visibility = visibility
            
        # Set featured status if provided
        featured = request.data.get('featured')
        if featured is not None:
            article.is_featured = featured.lower() == 'true'
            
        article.save()
        
        serializer = self.get_serializer(article)
        return Response(serializer.data)
    
    @action(detail=True, methods=['post'], permission_classes=[IsAdminUser])
    def reject(self, request, pk=None):
        """
        Endpoint for admins to reject an article
        """
        article = self.get_object()
        
        # Check if article is already approved
        if article.is_approved:
            return Response({
                'error': 'Cannot reject an already approved article'
            }, status=status.HTTP_400_BAD_REQUEST)
            
        # Check if article is already rejected
        if article.is_rejected:
            return Response({
                'error': 'Article is already rejected'
            }, status=status.HTTP_400_BAD_REQUEST)
            
        # Reject the article
        article.is_rejected = True
        article.rejected_by = request.user
        article.rejection_date = timezone.now()
        
        # Add rejection reason if provided
        rejection_reason = request.data.get('rejection_reason')
        if rejection_reason:
            article.rejection_reason = rejection_reason
        else:
            return Response({
                'error': 'Rejection reason is required'
            }, status=status.HTTP_400_BAD_REQUEST)
            
        article.save()
        
        serializer = self.get_serializer(article)
        return Response(serializer.data)
    
    @swagger_auto_schema(
        method='post',
        operation_description="Publish an approved article. Can only be done by the article author or an admin.",
        request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            properties={
                'visibility': openapi.Schema(type=openapi.TYPE_STRING, enum=['public', 'subscribers', 'private'], description='Optional visibility level for the article')
            }
        ),
        responses={
            200: ArticleSerializer,
            400: openapi.Response("Bad Request", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'error': openapi.Schema(type=openapi.TYPE_STRING)
                }
            )),
            403: openapi.Response("Forbidden", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'error': openapi.Schema(type=openapi.TYPE_STRING)
                }
            ))
        }
    )
    @action(detail=True, methods=['post'], permission_classes=[IsDoctorUser | IsAdminUser])
    def publish(self, request, pk=None):
        """
        Endpoint to publish an approved article
        """
        article = self.get_object()
        
        # Check if article is approved
        if not article.is_approved:
            return Response({
                'error': 'Article must be approved before publishing'
            }, status=status.HTTP_400_BAD_REQUEST)
            
        # Check if user is the author or admin
        is_admin = request.user.roles.filter(name='admin').exists()
        try:
            doctor = request.user.doctor
            is_author = (article.author == doctor)
            
            if not (is_author or is_admin):
                return Response({
                    'error': 'You can only publish your own articles'
                }, status=status.HTTP_403_FORBIDDEN)
                
            # Publish the article
            article.is_published = True
            article.publish_date = timezone.now()
            
            # Set visibility if provided
            visibility = request.data.get('visibility')
            if visibility in ['public', 'subscribers', 'private']:
                article.visibility = visibility
            
            article.save()
            
            serializer = self.get_serializer(article)
            return Response(serializer.data)
            
        except Doctor.DoesNotExist:
            if is_admin:
                # Admin can publish regardless
                article.is_published = True
                article.publish_date = timezone.now()
                
                # Set visibility if provided
                visibility = request.data.get('visibility')
                if visibility in ['public', 'subscribers', 'private']:
                    article.visibility = visibility
                    
                article.save()
                
                serializer = self.get_serializer(article)
                return Response(serializer.data)
            else:
                return Response({
                    'error': 'Doctor profile not found'
                }, status=status.HTTP_404_NOT_FOUND)
    
    @swagger_auto_schema(
        method='post',
        operation_description="Unpublish a published article. Can only be done by the article author or an admin.",
        responses={
            200: ArticleSerializer,
            400: openapi.Response("Bad Request", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'error': openapi.Schema(type=openapi.TYPE_STRING)
                }
            )),
            403: openapi.Response("Forbidden", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'error': openapi.Schema(type=openapi.TYPE_STRING)
                }
            )),
            404: openapi.Response("Not Found", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'error': openapi.Schema(type=openapi.TYPE_STRING)
                }
            ))
        }
    )
    @action(detail=True, methods=['post'], permission_classes=[IsDoctorUser | IsAdminUser])
    def unpublish(self, request, pk=None):
        """
        Endpoint to unpublish an article
        """
        article = self.get_object()
        
        # Check if article is published
        if not article.is_published:
            return Response({
                'error': 'Article is not published'
            }, status=status.HTTP_400_BAD_REQUEST)
            
        # Check if user is the author or admin
        is_admin = request.user.roles.filter(name='admin').exists()
        try:
            doctor = request.user.doctor
            is_author = (article.author == doctor)
            
            if not (is_author or is_admin):
                return Response({
                    'error': 'You can only unpublish your own articles'
                }, status=status.HTTP_403_FORBIDDEN)
                
            # Unpublish the article
            article.is_published = False
            article.save()
            
            serializer = self.get_serializer(article)
            return Response(serializer.data)
            
        except Doctor.DoesNotExist:
            if is_admin:
                # Admin can unpublish regardless
                article.is_published = False
                article.save()
                
                serializer = self.get_serializer(article)
                return Response(serializer.data)
            else:
                return Response({
                    'error': 'Doctor profile not found'
                }, status=status.HTTP_404_NOT_FOUND)
    
    @swagger_auto_schema(
        method='post',
        operation_description="Increment the view count for an article. Call this when a user views the article.",
        responses={
            200: openapi.Response("Success", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'status': openapi.Schema(type=openapi.TYPE_STRING),
                    'view_count': openapi.Schema(type=openapi.TYPE_INTEGER)
                }
            ))
        }
    )
    @action(detail=True, methods=['post'])
    def view(self, request, pk=None):
        """
        Endpoint to increment view count for an article
        """
        article = self.get_object()
        article.view_count += 1
        article.save()
        
        return Response({'status': 'view counted', 'view_count': article.view_count})
    
    @swagger_auto_schema(
        method='get',
        operation_description="Endpoint for doctors to view their own articles",
        manual_parameters=[
            openapi.Parameter('status', openapi.IN_QUERY, type=openapi.TYPE_STRING, enum=['draft', 'published', 'pending', 'approved'], 
                            description="Filter by article status"),
            openapi.Parameter('visibility', openapi.IN_QUERY, type=openapi.TYPE_STRING, enum=['public', 'subscribers', 'private'], 
                            description="Filter by visibility level"),
            openapi.Parameter('sort_by', openapi.IN_QUERY, type=openapi.TYPE_STRING, enum=["popular", "newest", "oldest"], 
                            description="Sort by view count (popular), newest first (newest), or oldest first (oldest)")
        ],
        responses={
            200: ArticleSerializer(many=True)
        }
    )
    @action(detail=False, methods=['get'])
    def my_articles(self, request):
        """
        Endpoint for doctors to view their own articles
        """
        try:
            doctor = request.user.doctor
            articles = Article.objects.filter(author=doctor)
            
            # Optional filtering
            status_param = request.query_params.get('status')
            if status_param == 'draft':
                articles = articles.filter(is_published=False)
            elif status_param == 'published':
                articles = articles.filter(is_published=True)
            elif status_param == 'pending':
                articles = articles.filter(is_approved=False)
            elif status_param == 'approved':
                articles = articles.filter(is_approved=True)
                
            # Filter by visibility
            visibility = request.query_params.get('visibility')
            if visibility:
                articles = articles.filter(visibility=visibility)
                
            # Sort options
            sort_by = self.request.query_params.get('sort_by')
            if sort_by:
                if sort_by == 'popular':
                    articles = articles.order_by('-view_count')
                elif sort_by == 'newest':
                    articles = articles.order_by('-created_at')
                elif sort_by == 'oldest':
                    articles = articles.order_by('created_at')
            
            serializer = self.get_serializer(articles, many=True)
            return Response(serializer.data)
        except (Doctor.DoesNotExist, AttributeError):
            return Response({
                'error': 'Doctor profile not found'
            }, status=status.HTTP_404_NOT_FOUND)
            
    @swagger_auto_schema(
        method='get',
        operation_description="Get featured articles that have been marked as featured by an admin",
        manual_parameters=[
            openapi.Parameter('limit', openapi.IN_QUERY, type=openapi.TYPE_INTEGER, 
                            description="Maximum number of featured articles to return (default: 5)")
        ],
        responses={
            200: ArticleSerializer(many=True)
        }
    )
    @action(detail=False, methods=['get'])
    def featured(self, request):
        """
        Endpoint to get featured articles
        """
        # Get the base queryset with proper permissions applied
        queryset = self.get_queryset()
        
        # Further filter to only featured articles
        queryset = queryset.filter(is_featured=True)
        
        # Limit to a reasonable number
        limit = request.query_params.get('limit', 5)
        try:
            limit = int(limit)
        except ValueError:
            limit = 5
            
        queryset = queryset[:limit]
        
        serializer = self.get_serializer(queryset, many=True)
        return Response(serializer.data)
        
    @swagger_auto_schema(
        method='get',
        operation_description="Get popular articles sorted by view count (most viewed first)",
        manual_parameters=[
            openapi.Parameter('limit', openapi.IN_QUERY, type=openapi.TYPE_INTEGER, 
                            description="Maximum number of popular articles to return (default: 10)")
        ],
        responses={
            200: ArticleSerializer(many=True)
        }
    )
    @action(detail=False, methods=['get'])
    def popular(self, request):
        """
        Endpoint to get popular articles based on view count
        """
        # Get the base queryset with proper permissions applied
        queryset = self.get_queryset()
        
        # Order by view count (most viewed first)
        queryset = queryset.order_by('-view_count')
        
        # Limit to a reasonable number
        limit = request.query_params.get('limit', 10)
        try:
            limit = int(limit)
        except ValueError:
            limit = 10
            
        queryset = queryset[:limit]
        
        serializer = self.get_serializer(queryset, many=True)
        return Response(serializer.data)
        
    @swagger_auto_schema(
        method='get',
        operation_description="Get recently published articles sorted by publish date (newest first)",
        manual_parameters=[
            openapi.Parameter('limit', openapi.IN_QUERY, type=openapi.TYPE_INTEGER, 
                            description="Maximum number of recent articles to return (default: 10)")
        ],
        responses={
            200: ArticleSerializer(many=True)
        }
    )
    @action(detail=False, methods=['get'])
    def recent(self, request):
        """
        Endpoint to get recently published articles
        """
        # Get the base queryset with proper permissions applied
        queryset = self.get_queryset()
        
        # Order by publish date (most recent first)
        queryset = queryset.filter(publish_date__isnull=False).order_by('-publish_date')
        
        # Limit to a reasonable number
        limit = request.query_params.get('limit', 10)
        try:
            limit = int(limit)
        except ValueError:
            limit = 10
            
        queryset = queryset[:limit]
        
        serializer = self.get_serializer(queryset, many=True)
        return Response(serializer.data)
        
    @swagger_auto_schema(
        method='get',
        operation_description="Get articles related to specific health conditions",
        manual_parameters=[
            openapi.Parameter('condition', openapi.IN_QUERY, type=openapi.TYPE_STRING, required=True, 
                            description="The health condition to find articles for (e.g., 'diabetes')"),
            openapi.Parameter('sort_by', openapi.IN_QUERY, type=openapi.TYPE_STRING, enum=["popular", "newest", "oldest"], 
                            description="Sort by popularity, newest, or oldest")
        ],
        responses={
            200: ArticleSerializer(many=True),
            400: openapi.Response("Bad Request", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'error': openapi.Schema(type=openapi.TYPE_STRING)
                }
            ))
        }
    )
    @action(detail=False, methods=['get'])
    def by_condition(self, request):
        """
        Endpoint to get articles related to specific health conditions
        """
        condition = request.query_params.get('condition')
        if not condition:
            return Response({
                'error': 'condition parameter is required'
            }, status=status.HTTP_400_BAD_REQUEST)
            
        # Get the base queryset with proper permissions applied
        queryset = self.get_queryset()
        
        # Filter by related conditions
        queryset = queryset.filter(related_conditions__icontains=condition)
        
        serializer = self.get_serializer(queryset, many=True)
        return Response(serializer.data)
    
    @action(detail=True, methods=['get'], permission_classes=[permissions.IsAuthenticated])
    def export_word(self, request, pk=None):
        """
        Export article to a Word document
        """
        article = self.get_object()
        
        # Check if user can view this article (same logic as get_queryset)
        if not article.is_published and not request.user.is_staff:
            # If not published, only author and admin can export
            if hasattr(request.user, 'doctor') and article.author.user != request.user:
                return Response({
                    'error': 'You do not have permission to export this article'
                }, status=status.HTTP_403_FORBIDDEN)
        
        # Create Word document
        doc = Document()
        
        # Add title
        title_paragraph = doc.add_heading(article.title, level=1)
        
        # Add metadata
        doc.add_paragraph(f"Author: {article.author.user.get_full_name()}")
        doc.add_paragraph(f"Category: {article.get_category_display()}")
        doc.add_paragraph(f"Published: {article.publish_date.strftime('%B %d, %Y') if article.publish_date else 'Not published'}")
        doc.add_paragraph(f"Reading Time: {article.reading_time} minutes")
        
        if article.tags:
            doc.add_paragraph(f"Tags: {article.tags}")
        
        if article.related_conditions:
            doc.add_paragraph(f"Related Conditions: {article.related_conditions}")
        
        # Add separator
        doc.add_paragraph("─" * 50)
        
        # Add summary if exists
        if article.summary:
            doc.add_heading("Summary", level=2)
            doc.add_paragraph(article.summary)
        
        # Add main content
        doc.add_heading("Content", level=2)
        
        # Split content into paragraphs and add them
        content_paragraphs = article.content.split('\n')
        for paragraph_text in content_paragraphs:
            if paragraph_text.strip():
                doc.add_paragraph(paragraph_text.strip())
        
        # Add footer with export info
        doc.add_paragraph("")
        doc.add_paragraph("─" * 50)
        doc.add_paragraph(f"Exported on: {timezone.now().strftime('%B %d, %Y at %I:%M %p')}")
        doc.add_paragraph(f"Exported by: {request.user.get_full_name()}")
        
        # Save document to memory
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        # Create response
        response = HttpResponse(
            doc_io.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        # Set filename - sanitize title for filename
        safe_title = "".join(c for c in article.title if c.isalnum() or c in (' ', '-', '_')).rstrip()
        safe_title = safe_title.replace(' ', '_')[:50]  # Limit length
        filename = f"article_{safe_title}.docx"
        
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
 
 
class ArticleCommentViewSet(viewsets.ModelViewSet):
    queryset = ArticleComment.objects.all()
    serializer_class = ArticleCommentSerializer
    permission_classes = [permissions.IsAuthenticated]
    
    @swagger_auto_schema(
        operation_description="List all comments with filtering options",
        manual_parameters=[
            openapi.Parameter('article_id', openapi.IN_QUERY, type=openapi.TYPE_STRING, 
                            description="Filter by article ID"),
            openapi.Parameter('user_id', openapi.IN_QUERY, type=openapi.TYPE_STRING, 
                            description="Filter by user ID"),
            openapi.Parameter('top_level_only', openapi.IN_QUERY, type=openapi.TYPE_BOOLEAN, 
                            description="Only return top-level comments (no replies)")
        ],
        responses={
            200: ArticleCommentSerializer(many=True)
        }
    )
    def list(self, request, *args, **kwargs):
        return super().list(request, *args, **kwargs)
    
    def get_queryset(self):
        queryset = ArticleComment.objects.all()
        
        # Only show top-level comments by default
        top_level_only = self.request.query_params.get('top_level_only')
        if top_level_only and top_level_only.lower() == 'true':
            queryset = queryset.filter(parent_comment__isnull=True)
            
        # Filter by article
        article_id = self.request.query_params.get('article_id')
        if article_id:
            queryset = queryset.filter(article_id=article_id)
            
        # Filter by user
        user_id = self.request.query_params.get('user_id')
        if user_id:
            queryset = queryset.filter(user_id=user_id)
            
        return queryset
    
    def perform_create(self, serializer):
        serializer.save(user=self.request.user)
    
    def perform_update(self, serializer):
        # Ensure users can only edit their own comments
        comment = self.get_object()
        if comment.user != self.request.user and not (self.request.user.is_authenticated and self.request.user.roles.filter(name='admin').exists()):
            raise serializers.ValidationError("You can only edit your own comments")
        serializer.save()
    
    @swagger_auto_schema(
        method='post',
        operation_description="Like a comment. Users can only like a comment once.",
        responses={
            200: openapi.Response("Success", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'status': openapi.Schema(type=openapi.TYPE_STRING),
                    'like_count': openapi.Schema(type=openapi.TYPE_INTEGER)
                }
            )),
            400: openapi.Response("Bad Request", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'error': openapi.Schema(type=openapi.TYPE_STRING)
                }
            ))
        }
    )
    @action(detail=True, methods=['post'])
    def like(self, request, pk=None):
        """
        Endpoint to like a comment
        """
        comment = self.get_object()
        user = request.user
        
        # Check if user already liked this comment
        existing_like = ArticleCommentLike.objects.filter(comment=comment, user=user).exists()
        if existing_like:
            return Response({
                'error': 'You have already liked this comment'
            }, status=status.HTTP_400_BAD_REQUEST)
            
        # Create a new like
        ArticleCommentLike.objects.create(comment=comment, user=user)
        
        # Update the comment's like count
        comment.like_count += 1
        comment.save()
        
        return Response({'status': 'comment liked', 'like_count': comment.like_count})
    
    @swagger_auto_schema(
        method='post',
        operation_description="Remove a like from a comment. Users can only unlike comments they've previously liked.",
        responses={
            200: openapi.Response("Success", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'status': openapi.Schema(type=openapi.TYPE_STRING),
                    'like_count': openapi.Schema(type=openapi.TYPE_INTEGER)
                }
            )),
            400: openapi.Response("Bad Request", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'error': openapi.Schema(type=openapi.TYPE_STRING)
                }
            ))
        }
    )
    @action(detail=True, methods=['post'])
    def unlike(self, request, pk=None):
        """
        Endpoint to unlike a comment
        """
        comment = self.get_object()
        user = request.user
        
        # Check if user has liked this comment
        try:
            like = ArticleCommentLike.objects.get(comment=comment, user=user)
            like.delete()
            
            # Update the comment's like count
            comment.like_count = max(0, comment.like_count - 1)
            comment.save()
            
            return Response({'status': 'comment unliked', 'like_count': comment.like_count})
        except ArticleCommentLike.DoesNotExist:
            return Response({
                'error': 'You have not liked this comment'
            }, status=status.HTTP_400_BAD_REQUEST)
    
    @swagger_auto_schema(
        method='post',
        operation_description="Reply to a top-level comment. Only one level of nesting is allowed (no replies to replies).",
        request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            required=['content'],
            properties={
                'content': openapi.Schema(type=openapi.TYPE_STRING, description="The text content of the reply")
            }
        ),
        responses={
            200: ArticleCommentReplySerializer,
            400: openapi.Response("Bad Request", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'error': openapi.Schema(type=openapi.TYPE_STRING)
                }
            ))
        }
    )
    @action(detail=True, methods=['post'])
    def reply(self, request, pk=None):
        """
        Endpoint to reply to a comment
        """
        parent_comment = self.get_object()
        
        # Ensure this is a top-level comment (no nested replies beyond 1 level)
        if parent_comment.parent_comment is not None:
            return Response({
                'error': 'Cannot reply to a reply. Only one level of nesting is allowed.'
            }, status=status.HTTP_400_BAD_REQUEST)
            
        # Get the content from request data
        content = request.data.get('content')
        if not content:
            return Response({
                'error': 'Content is required for a reply'
            }, status=status.HTTP_400_BAD_REQUEST)
            
        # Create the reply
        user = request.user
        is_doctor = user.roles.filter(name='doctor').exists()
        
        reply = ArticleComment.objects.create(
            article=parent_comment.article,
            content=content,
            user=user,
            is_doctor=is_doctor,
            parent_comment=parent_comment
        )
        
        serializer = ArticleCommentReplySerializer(reply)
        return Response(serializer.data)


class PackageViewSet(viewsets.ModelViewSet):
    """
    ViewSet for managing subscription packages
    """
    queryset = Package.objects.all()
    serializer_class = PackageSerializer
    permission_classes = [permissions.IsAuthenticated]
    
    def get_permissions(self):
        """
        Admin can create/update/delete packages
        All authenticated users can view packages
        """
        if self.action in ['create', 'update', 'partial_update', 'destroy']:
            permission_classes = [IsAdminUser]
        else:
            permission_classes = [permissions.IsAuthenticated]
        return [permission() for permission in permission_classes]
    
    def get_queryset(self):
        """
        Filter packages - only show active packages to non-admin users
        """
        queryset = Package.objects.all()
        
        # Non-admin users only see active packages
        if not (self.request.user.is_authenticated and self.request.user.roles.filter(name='admin').exists()):
            queryset = queryset.filter(is_active=True)
        
        return queryset.order_by('price')


class PatientSubscriptionViewSet(viewsets.ModelViewSet):
    """
    ViewSet for managing patient subscriptions
    """
    queryset = PatientSubscription.objects.all()
    serializer_class = PatientSubscriptionSerializer
    permission_classes = [permissions.IsAuthenticated]
    
    def get_permissions(self):
        """
        Admin can create/view/update all subscriptions
        Patients can only view their own subscriptions
        """
        if self.action in ['create', 'update', 'partial_update', 'destroy']:
            permission_classes = [IsAdminUser]
        else:
            permission_classes = [permissions.IsAuthenticated]
        return [permission() for permission in permission_classes]
    
    def get_queryset(self):
        """
        Filter subscriptions by user role
        """
        queryset = PatientSubscription.objects.all()
        
        if self.request.user.is_authenticated and self.request.user.roles.filter(name='admin').exists():
            # Admin can see all subscriptions
            pass
        elif self.request.user.is_authenticated and self.request.user.roles.filter(name='patient').exists():
            # Patients can only see their own subscriptions
            try:
                patient = self.request.user.patient
                queryset = queryset.filter(patient=patient)
            except Patient.DoesNotExist:
                return PatientSubscription.objects.none()
        else:
            return PatientSubscription.objects.none()
        
        return queryset.order_by('-created_at')
    
    @action(detail=False, methods=['post'], permission_classes=[IsPatientUser])
    def subscribe(self, request):
        """
        Create a new subscription with payment processing
        """
        try:
            patient = request.user.patient
            package_id = request.data.get('package_id')
            payment_method = request.data.get('payment_method', 'pesapal')
            subscription_frequency = request.data.get('subscription_frequency', 'MONTHLY')
            
            if not package_id:
                return Response({
                    'error': 'Package ID is required'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            try:
                package = Package.objects.get(id=package_id, is_active=True)
            except Package.DoesNotExist:
                return Response({
                    'error': 'Package not found or inactive'
                }, status=status.HTTP_404_NOT_FOUND)
            
            # Check if patient has an active subscription
            active_subscription = PatientSubscription.objects.filter(
                patient=patient,
                status='active',
                end_date__gte=timezone.now().date()
            ).first()
            
            if active_subscription:
                return Response({
                    'error': 'Patient already has an active subscription',
                    'active_subscription': PatientSubscriptionSerializer(active_subscription).data
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Create payment record
            import secrets
            payment_reference = f"PAY_{secrets.token_hex(8).upper()}"
            
            payment = Payment.objects.create(
                reference=payment_reference,
                amount=package.price,
                payment_method=payment_method,
                status='pending'
            )
            
            # Create subscription
            start_date = timezone.now().date()
            end_date = start_date + timedelta(days=package.duration_days)
            
            subscription = PatientSubscription.objects.create(
                patient=patient,
                package=package,
                payment=payment,
                start_date=start_date,
                end_date=end_date,
                status='pending'
            )
            
            # Return subscription details with payment info
            serializer = PatientSubscriptionSerializer(subscription)
            return Response({
                'subscription': serializer.data,
                'payment_reference': payment_reference,
                'payment_id': payment.id,
                'payment_url': f'/api/payments/{payment.id}/process/',
                'message': 'Subscription created. Please complete payment.'
            }, status=status.HTTP_201_CREATED)
            
        except Patient.DoesNotExist:
            return Response({
                'error': 'Patient profile not found'
            }, status=status.HTTP_404_NOT_FOUND)
    
    @action(detail=True, methods=['post'], permission_classes=[IsPatientUser])
    def cancel(self, request, pk=None):
        """
        Cancel an active subscription
        """
        try:
            subscription = self.get_object()
            
            # Check if subscription belongs to current user
            if subscription.patient != request.user.patient:
                return Response({
                    'error': 'Subscription not found'
                }, status=status.HTTP_404_NOT_FOUND)
            
            if subscription.status != 'active':
                return Response({
                    'error': 'Only active subscriptions can be cancelled'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Cancel subscription
            subscription.status = 'cancelled'
            subscription.save()
            
            serializer = PatientSubscriptionSerializer(subscription)
            return Response({
                'subscription': serializer.data,
                'message': 'Subscription cancelled successfully'
            }, status=status.HTTP_200_OK)
            
        except PatientSubscription.DoesNotExist:
            return Response({
                'error': 'Subscription not found'
            }, status=status.HTTP_404_NOT_FOUND)
        except Patient.DoesNotExist:
            return Response({
                'error': 'Patient profile not found'
            }, status=status.HTTP_404_NOT_FOUND)
    
    @action(detail=False, methods=['get'], permission_classes=[IsPatientUser])
    def active(self, request):
        """
        Get current active subscription for the patient
        """
        try:
            patient = request.user.patient
            
            active_subscription = PatientSubscription.objects.filter(
                patient=patient,
                status='active',
                end_date__gte=timezone.now().date()
            ).first()
            
            if not active_subscription:
                return Response({
                    'message': 'No active subscription found'
                }, status=status.HTTP_404_NOT_FOUND)
            
            serializer = PatientSubscriptionSerializer(active_subscription)
            return Response({
                'subscription': serializer.data
            }, status=status.HTTP_200_OK)
            
        except Patient.DoesNotExist:
            return Response({
                'error': 'Patient profile not found'
            }, status=status.HTTP_404_NOT_FOUND)
    
    @action(detail=False, methods=['post'], permission_classes=[IsPatientUser])
    def upgrade(self, request):
        """
        Upgrade subscription to a higher-tier package
        """
        try:
            patient = request.user.patient
            new_package_id = request.data.get('package_id')
            
            if not new_package_id:
                return Response({
                    'error': 'Package ID is required'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            subscription_manager = SubscriptionManager()
            result = subscription_manager.upgrade_subscription(patient, new_package_id)
            
            if result['success']:
                return Response({
                    'payment_id': result['payment_id'],
                    'payment_reference': result['payment_reference'],
                    'prorated_amount': result['prorated_amount'],
                    'remaining_days': result['remaining_days'],
                    'payment_url': f'/api/payments/{result["payment_id"]}/process/',
                    'message': result['message']
                }, status=status.HTTP_200_OK)
            else:
                return Response({
                    'error': result['error']
                }, status=status.HTTP_400_BAD_REQUEST)
                
        except Patient.DoesNotExist:
            return Response({
                'error': 'Patient profile not found'
            }, status=status.HTTP_404_NOT_FOUND)
    
    @action(detail=False, methods=['post'], permission_classes=[IsPatientUser])
    def downgrade(self, request):
        """
        Downgrade subscription to a lower-tier package
        """
        try:
            patient = request.user.patient
            new_package_id = request.data.get('package_id')
            
            if not new_package_id:
                return Response({
                    'error': 'Package ID is required'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            subscription_manager = SubscriptionManager()
            result = subscription_manager.downgrade_subscription(patient, new_package_id)
            
            if result['success']:
                return Response({
                    'effective_date': result['effective_date'],
                    'credit_amount': result.get('credit_amount', 0),
                    'message': result['message']
                }, status=status.HTTP_200_OK)
            else:
                return Response({
                    'error': result['error']
                }, status=status.HTTP_400_BAD_REQUEST)
                
        except Patient.DoesNotExist:
            return Response({
                'error': 'Patient profile not found'
            }, status=status.HTTP_404_NOT_FOUND)
    
    @action(detail=True, methods=['post'], permission_classes=[IsPatientUser])
    def renew(self, request, pk=None):
        """
        Renew a subscription for another period
        """
        try:
            subscription = self.get_object()
            
            # Check if subscription belongs to current user
            if subscription.patient != request.user.patient:
                return Response({
                    'error': 'Subscription not found'
                }, status=status.HTTP_404_NOT_FOUND)
            
            subscription_manager = SubscriptionManager()
            result = subscription_manager.renew_subscription(subscription)
            
            if result['success']:
                return Response({
                    'payment_id': result['payment_id'],
                    'payment_reference': result['payment_reference'],
                    'new_end_date': result['new_end_date'],
                    'payment_url': f'/api/payments/{result["payment_id"]}/process/',
                    'message': result['message']
                }, status=status.HTTP_200_OK)
            else:
                return Response({
                    'error': result['error']
                }, status=status.HTTP_400_BAD_REQUEST)
                
        except PatientSubscription.DoesNotExist:
            return Response({
                'error': 'Subscription not found'
            }, status=status.HTTP_404_NOT_FOUND)
        except Patient.DoesNotExist:
            return Response({
                'error': 'Patient profile not found'
            }, status=status.HTTP_404_NOT_FOUND)
    
    @action(detail=True, methods=['get'], permission_classes=[IsPatientUser])
    def usage(self, request, pk=None):
        """
        Get subscription usage statistics
        """
        try:
            subscription = self.get_object()
            
            # Check if subscription belongs to current user
            if subscription.patient != request.user.patient:
                return Response({
                    'error': 'Subscription not found'
                }, status=status.HTTP_404_NOT_FOUND)
            
            subscription_manager = SubscriptionManager()
            usage_stats = subscription_manager.get_subscription_usage(subscription)
            
            if 'error' in usage_stats:
                return Response({
                    'error': usage_stats['error']
                }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            
            return Response(usage_stats, status=status.HTTP_200_OK)
                
        except PatientSubscription.DoesNotExist:
            return Response({
                'error': 'Subscription not found'
            }, status=status.HTTP_404_NOT_FOUND)
        except Patient.DoesNotExist:
            return Response({
                'error': 'Patient profile not found'
            }, status=status.HTTP_404_NOT_FOUND)
    
    @action(detail=True, methods=['post'], permission_classes=[IsPatientUser])
    def cancel_subscription(self, request, pk=None):
        """
        Cancel an active subscription
        """
        try:
            subscription = self.get_object()
            patient = request.user.patient
            
            if subscription.patient != patient:
                return Response({
                    'error': 'You can only cancel your own subscriptions'
                }, status=status.HTTP_403_FORBIDDEN)
            
            if subscription.status != 'active':
                return Response({
                    'error': 'Only active subscriptions can be cancelled'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            subscription.status = 'cancelled'
            subscription.save()
            
            return Response({
                'message': 'Subscription cancelled successfully'
            }, status=status.HTTP_200_OK)
            
        except Patient.DoesNotExist:
            return Response({
                'error': 'Patient profile not found'
            }, status=status.HTTP_404_NOT_FOUND)


class PaymentViewSet(viewsets.ModelViewSet):
    """
    ViewSet for handling payment processing with Pesapal integration
    """
    queryset = Payment.objects.all()
    serializer_class = PaymentSerializer
    permission_classes = [permissions.IsAuthenticated]
    
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.pesapal_client = PesapalClient()
    
    def get_permissions(self):
        """
        Only authenticated users can access payment endpoints
        """
        if self.action in ['create', 'update', 'partial_update', 'destroy']:
            permission_classes = [IsAdminUser]
        elif self.action in ['callback', 'ipn']:
            permission_classes = [permissions.AllowAny]
        else:
            permission_classes = [permissions.IsAuthenticated]
        return [permission() for permission in permission_classes]
    
    def get_queryset(self):
        """
        Filter payments by user role
        """
        queryset = Payment.objects.all()
        
        if self.request.user.is_authenticated and self.request.user.roles.filter(name='admin').exists():
            # Admin can see all payments
            pass
        elif self.request.user.is_authenticated and self.request.user.roles.filter(name='patient').exists():
            # Patients can only see payments for their subscriptions
            try:
                patient = self.request.user.patient
                queryset = queryset.filter(subscriptions__patient=patient)
            except Patient.DoesNotExist:
                return Payment.objects.none()
        else:
            return Payment.objects.none()
        
        return queryset.order_by('-created_at')
    
    @action(detail=True, methods=['post'], permission_classes=[permissions.IsAuthenticated])
    def process(self, request, pk=None):
        """
        Process payment using Pesapal API
        """
        try:
            payment = self.get_object()
            
            # Check if payment belongs to current user (if not admin)
            if not request.user.roles.filter(name='admin').exists():
                try:
                    patient = request.user.patient
                    if not payment.subscriptions.filter(patient=patient).exists():
                        return Response({
                            'error': 'Payment not found'
                        }, status=status.HTTP_404_NOT_FOUND)
                except Patient.DoesNotExist:
                    return Response({
                        'error': 'Patient profile not found'
                    }, status=status.HTTP_404_NOT_FOUND)
                            
            if payment.status != 'pending':
                return Response({
                    'error': 'Payment already processed or failed'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Get subscription details
            subscription = payment.subscriptions.first()
            if not subscription:
                return Response({
                    'error': 'No subscription associated with this payment'
                }, status=status.HTTP_400_BAD_REQUEST)
            

            
            # Update payment status to processing
            payment.status = 'processing'
            payment.save()
            
            # Prepare billing address from patient and user data
            patient = subscription.patient
            user = patient.user
            
            # Create order data for Pesapal
            order_data = {
                "id": f"SUB_{payment.id}",
                "currency": payment.currency,
                "amount": float(payment.amount),
                "description": f"Healthcare subscription - {subscription.package.name}",
                "callback_url": f"{settings.FRONTEND_URL}/payment/callback?payment_id={payment.id}",
                "notification_id": getattr(settings, 'PESAPAL_IPN_ID', ''),
                "account_number": f"PAT_{patient.id}",
                "billing_address": {
                    "email_address": user.email,
                    "phone_number": '254795941990',
                    "country_code": "KE",
                    "first_name": user.first_name,
                    "last_name": user.last_name,
                    "line_1": "Moi Avenue",      
                    "line_2": "Suite 12",          
                    "city": "Nairobi",           
                    "state": "Nairobi County",    
                    "postal_code": "00100",
                    "zip_code": "00100" 
                }
            }            

            pesapal_response = self.pesapal_client.submit_order_request(order_data)
            
            if pesapal_response.get("error"):
                payment.status = 'failed'
                payment.gateway_response = pesapal_response
                payment.save()

                
                return Response({
                    'error': 'Payment processing failed',
                    'details': pesapal_response.get('error', {})
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Update payment with Pesapal response
            payment.gateway_transaction_id = pesapal_response.get('order_tracking_id')
            payment.gateway_response = pesapal_response
            payment.save()
            
            return Response({
                'payment_id': payment.id,
                'payment_reference': payment.reference,
                'amount': payment.amount,
                'currency': payment.currency,
                'payment_method': payment.payment_method,
                'redirect_url': pesapal_response.get('redirect_url'),
                'order_tracking_id': pesapal_response.get('order_tracking_id'),
                'message': 'Payment processing initiated. You will be redirected to Pesapal.'
            }, status=status.HTTP_200_OK)
            
        except Payment.DoesNotExist:
            return Response({
                'error': 'Payment not found'
            }, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            return Response({
                'error': 'Payment processing failed',
                'details': str(e)
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
    @action(detail=True, methods=['post'], permission_classes=[permissions.AllowAny])
    def callback(self, request, pk=None):
        """
        Handle payment callback from Pesapal
        """
        try:
            payment = Payment.objects.get(id=pk)
            
            # Get callback data from Pesapal
            order_tracking_id = request.data.get('OrderTrackingId') or request.GET.get('OrderTrackingId')
            merchant_reference = request.data.get('OrderMerchantReference') or request.GET.get('OrderMerchantReference')
            
            if not order_tracking_id:
                return Response({
                    'error': 'Order tracking ID is required'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Verify payment status with Pesapal
            status_response = self.pesapal_client.get_transaction_status(order_tracking_id)
            
            if "error" in status_response:
                return Response({
                    'error': 'Failed to verify payment status',
                    'details': status_response.get('error', {})
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Update payment record
            payment.gateway_transaction_id = order_tracking_id
            payment.gateway_response = status_response
            
            payment_status = status_response.get('payment_status_description', '').upper()
            
            if payment_status == 'COMPLETED':
                payment.status = 'completed'
                payment.save()
                
                # Activate associated subscriptions
                subscriptions = payment.subscriptions.all()
                for subscription in subscriptions:
                    subscription.status = 'active'
                    subscription.save()
                
                return Response({
                    'message': 'Payment completed successfully',
                    'payment_status': 'completed',
                    'order_tracking_id': order_tracking_id
                }, status=status.HTTP_200_OK)
            
            elif payment_status in ['FAILED', 'INVALID']:
                payment.status = 'failed'
                payment.save()
                
                return Response({
                    'message': 'Payment failed',
                    'payment_status': 'failed',
                    'order_tracking_id': order_tracking_id
                }, status=status.HTTP_200_OK)
            
            else:
                payment.status = 'pending'
                payment.save()
                
                return Response({
                    'message': 'Payment status pending',
                    'payment_status': 'pending',
                    'order_tracking_id': order_tracking_id
                }, status=status.HTTP_200_OK)
                
        except Payment.DoesNotExist:
            return Response({
                'error': 'Payment not found'
            }, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            return Response({
                'error': 'Payment callback processing failed',
                'details': str(e)
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
    @action(detail=False, methods=['post'], permission_classes=[permissions.AllowAny])
    def ipn(self, request):
        """
        Handle Instant Payment Notification (IPN) from Pesapal
        """
        try:
            # Get IPN data from Pesapal
            order_tracking_id = request.data.get('OrderTrackingId')
            merchant_reference = request.data.get('OrderMerchantReference')
            notification_type = request.data.get('OrderNotificationType')
            
            if not order_tracking_id:
                return Response({
                    'error': 'Order tracking ID is required'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Find payment by order tracking ID
            try:
                payment = Payment.objects.get(gateway_transaction_id=order_tracking_id)
            except Payment.DoesNotExist:
                return Response({
                    'error': 'Payment not found'
                }, status=status.HTTP_404_NOT_FOUND)
            
            # Verify payment status with Pesapal
            status_response = self.pesapal_client.get_transaction_status(order_tracking_id)
            
            if "error" in status_response:
                return Response({
                    'error': 'Failed to verify payment status'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Update payment record - preserve existing gateway_response data
            existing_response = payment.gateway_response or {}
            upgrade_cancel_subscription_id = existing_response.get('upgrade_cancel_subscription_id')
            
            # Merge with new status response
            payment.gateway_response = {**existing_response, **status_response}
            payment_status = status_response.get('payment_status_description', '').upper()
            
            if payment_status == 'COMPLETED':
                payment.status = 'completed'
                payment.save()
                
                # Check if this is an upgrade payment and handle old subscription cancellation
                if upgrade_cancel_subscription_id:
                    try:
                        old_subscription = PatientSubscription.objects.get(id=upgrade_cancel_subscription_id)
                        old_subscription.status = 'cancelled'
                        old_subscription.save()
                        
                        # Log the cancellation for tracking
                        import logging
                        logger = logging.getLogger(__name__)
                        logger.info(f"Cancelled old subscription {upgrade_cancel_subscription_id} after upgrade payment completion")
                    except PatientSubscription.DoesNotExist:
                        # Log error but don't fail the IPN processing
                        import logging
                        logger = logging.getLogger(__name__)
                        logger.error(f"Old subscription {upgrade_cancel_subscription_id} not found for cancellation after upgrade")
                
                # Activate associated subscriptions
                subscriptions = payment.subscriptions.all()
                for subscription in subscriptions:
                    subscription.status = 'active'
                    subscription.save()
                
                # Handle recurring payments
                if notification_type == 'RECURRING':
                    # Create new payment record for recurring payment
                    original_subscription = subscriptions.first()
                    if original_subscription:
                        # Calculate next payment date based on subscription frequency
                        next_payment_date = original_subscription.end_date + timedelta(days=1)
                        
                        # Create new payment record
                        import secrets
                        new_payment = Payment.objects.create(
                            reference=f"REC_{secrets.token_hex(8).upper()}",
                            amount=payment.amount,
                            currency=payment.currency,
                            payment_method=payment.payment_method,
                            status='completed',
                            gateway_transaction_id=order_tracking_id,
                            gateway_response=status_response
                        )
                        
                        # Extend subscription
                        original_subscription.end_date = next_payment_date + timedelta(days=original_subscription.package.duration_days)
                        original_subscription.save()
                        
                        # Link payment to subscription
                        new_payment.subscriptions.add(original_subscription)
                
            elif payment_status in ['FAILED', 'INVALID']:
                payment.status = 'failed'
                payment.save()
            
            return Response({
                'message': 'IPN processed successfully',
                'order_tracking_id': order_tracking_id
            }, status=status.HTTP_200_OK)
            
        except Exception as e:
            return Response({
                'error': 'IPN processing failed',
                'details': str(e)
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
    @action(detail=True, methods=['get'], permission_classes=[permissions.IsAuthenticated])
    def status(self, request, pk=None):
        """
        Check payment status and sync with Pesapal
        """
        try:
            payment = self.get_object()
            
            # Check if payment belongs to current user (if not admin)
            if not request.user.roles.filter(name='admin').exists():
                try:
                    patient = request.user.patient
                    if not payment.subscriptions.filter(patient=patient).exists():
                        return Response({
                            'error': 'Payment not found'
                        }, status=status.HTTP_404_NOT_FOUND)
                except Patient.DoesNotExist:
                    return Response({
                        'error': 'Patient profile not found'
                    }, status=status.HTTP_404_NOT_FOUND)
            
            # If payment is not completed and has a gateway transaction ID, sync with Pesapal
            if payment.status != 'completed' and payment.gateway_transaction_id:
                try:
                    status_response = self.pesapal_client.get_transaction_status(payment.gateway_transaction_id)
                    
                    if "error" not in status_response:
                        payment.gateway_response = status_response
                        payment_status = status_response.get('payment_status_description', '').upper()
                        
                        if payment_status == 'COMPLETED':
                            payment.status = 'completed'
                            payment.save()
                            
                            # Activate associated subscriptions
                            subscriptions = payment.subscriptions.all()
                            for subscription in subscriptions:
                                subscription.status = 'active'
                                subscription.save()
                                
                        elif payment_status in ['FAILED', 'INVALID']:
                            payment.status = 'failed'
                            payment.save()
                except Exception as e:
                    pass  # Continue with existing status if sync fails
            
            serializer = PaymentSerializer(payment)
            return Response(serializer.data, status=status.HTTP_200_OK)
            
        except Payment.DoesNotExist:
            return Response({
                'error': 'Payment not found'
            }, status=status.HTTP_404_NOT_FOUND)


class IsDoctorOrAdminUser(permissions.BasePermission):
    """
    Permission class to check if the user has doctor role OR admin role
    """
    def has_permission(self, request, view):
        # Check if user is authenticated
        if not request.user.is_authenticated:
            return False
        
        # Check if user has doctor role OR admin role
        return (request.user.roles.filter(name='doctor').exists() or 
                request.user.roles.filter(name='admin').exists())

class DoctorAvailabilityViewSet(viewsets.ModelViewSet):
    """
    ViewSet for managing doctor availability
    """
    queryset = DoctorAvailability.objects.all()
    serializer_class = DoctorAvailabilitySerializer
    permission_classes = [permissions.IsAuthenticated]
    
    def get_permissions(self):
        """
        Doctors can manage their own availability
        Admin can manage all availability
        All authenticated users can view availability
        """
        if self.action in ['create', 'update', 'partial_update', 'destroy']:
            permission_classes = [IsDoctorOrAdminUser]
        else:
            permission_classes = [permissions.IsAuthenticated]
        return [permission() for permission in permission_classes]
    
    def get_queryset(self):
        """
        Filter availability by doctor
        """
        queryset = DoctorAvailability.objects.all()
        
        # Filter by doctor_id if provided
        doctor_id = self.kwargs.get('doctor_id')

        print("This is the doctor id:", doctor_id)
        if doctor_id:
            queryset = queryset.filter(doctor=doctor_id, is_available=True)
        
        return queryset.order_by('weekday', 'start_time')
    
    def perform_create(self, serializer):
        """
        Ensure doctors can only create their own availability
        """
        if self.request.user.roles.filter(name='doctor').exists():
            try:
                doctor = self.request.user.doctor
                serializer.save(doctor=doctor)
            except Doctor.DoesNotExist:
                raise serializers.ValidationError("Doctor profile not found")
        else:
            # Admin can create for any doctor
            serializer.save()
    
    def perform_update(self, serializer):
        """
        Ensure doctors can only update their own availability
        """
        if self.request.user.roles.filter(name='doctor').exists():
            try:
                doctor = self.request.user.doctor
                if serializer.instance.doctor != doctor:
                    raise serializers.ValidationError("You can only update your own availability")
                serializer.save()
            except Doctor.DoesNotExist:
                raise serializers.ValidationError("Doctor profile not found")
        else:
            # Admin can update any availability
            serializer.save()
    
    def perform_destroy(self, instance):
        """
        Ensure doctors can only delete their own availability
        """
        if self.request.user.roles.filter(name='doctor').exists():
            try:
                doctor = self.request.user.doctor
                if instance.doctor != doctor:
                    raise serializers.ValidationError("You can only delete your own availability")
                instance.delete()
            except Doctor.DoesNotExist:
                raise serializers.ValidationError("Doctor profile not found")
        else:
            # Admin can delete any availability
            instance.delete()
    
    @action(detail=False, methods=['get'])
    def my_availability(self, request):
        """
        Get the current doctor's availability
        """
        if not request.user.roles.filter(name='doctor').exists():
            return Response({
                'error': 'Only doctors can access this endpoint'
            }, status=status.HTTP_403_FORBIDDEN)
        
        try:
            doctor = request.user.doctor
            availability = DoctorAvailability.objects.filter(doctor=doctor)
            serializer = self.get_serializer(availability, many=True)
            return Response(serializer.data)
        except Doctor.DoesNotExist:
            return Response({
                'error': 'Doctor profile not found'
            }, status=status.HTTP_404_NOT_FOUND)


class PackagePaymentTrackerViewSet(viewsets.ViewSet):
    """
    Package Payment Tracker API endpoints
    """
    permission_classes = [IsAdminUser]
    
    @swagger_auto_schema(
        operation_description="Get package payment tracker dashboard metrics",
        responses={
            200: openapi.Response("Success", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'total_active_subscribers': openapi.Schema(type=openapi.TYPE_INTEGER),
                    'expiring_this_week': openapi.Schema(type=openapi.TYPE_INTEGER),
                    'renewed_this_month': openapi.Schema(type=openapi.TYPE_INTEGER),
                    'overdue_not_renewed': openapi.Schema(type=openapi.TYPE_INTEGER),
                }
            ))
        }
    )
    @action(detail=False, methods=['get'], url_path='dashboard')
    def get_dashboard_metrics(self, request):
        """
        Get dashboard metrics for package payment tracker
        """
        from django.db.models import Count, Q
        from datetime import date, timedelta
        
        today = date.today()
        week_end = today + timedelta(days=7)
        month_start = today.replace(day=1)
        
        # Total Active Subscribers
        total_active = PatientSubscription.objects.filter(
            status='active',
            start_date__lte=today,
            end_date__gte=today
        ).count()
        
        # Expiring This Week
        expiring_week = PatientSubscription.objects.filter(
            status='active',
            end_date__range=[today, week_end]
        ).count()
        
        # Renewed This Month (new subscriptions this month)
        renewed_month = PatientSubscription.objects.filter(
            status='active',
            start_date__gte=month_start,
            start_date__lte=today
        ).count()
        
        # Overdue (expired but not renewed)
        overdue = PatientSubscription.objects.filter(
            status='expired',
            end_date__lt=today
        ).count()
        
        return Response({
            'total_active_subscribers': total_active,
            'expiring_this_week': expiring_week,
            'renewed_this_month': renewed_month,
            'overdue_not_renewed': overdue,
        })
    
    @swagger_auto_schema(
        operation_description="Get paginated list of patient subscriptions with filtering",
        manual_parameters=[
            openapi.Parameter('package_type', openapi.IN_QUERY, type=openapi.TYPE_STRING, 
                            description="Filter by package name"),
            openapi.Parameter('status', openapi.IN_QUERY, type=openapi.TYPE_STRING,
                            description="Filter by subscription status"),
            openapi.Parameter('search', openapi.IN_QUERY, type=openapi.TYPE_STRING,
                            description="Search by patient name or phone"),
            openapi.Parameter('page', openapi.IN_QUERY, type=openapi.TYPE_INTEGER,
                            description="Page number for pagination"),
            openapi.Parameter('page_size', openapi.IN_QUERY, type=openapi.TYPE_INTEGER,
                            description="Number of items per page"),
        ]
    )
    @action(detail=False, methods=['get'], url_path='subscriptions')
    def get_subscriptions(self, request):
        """
        Get paginated list of patient subscriptions with filtering and search
        """
        from django.core.paginator import Paginator
        from django.db.models import Q
        
        # Get query parameters
        package_type = request.query_params.get('package_type')
        status_filter = request.query_params.get('status')
        search_query = request.query_params.get('search')
        page = int(request.query_params.get('page', 1))
        page_size = int(request.query_params.get('page_size', 10))
        
        # Start with base queryset
        queryset = PatientSubscription.objects.select_related(
            'patient__user', 'package', 'payment'
        ).all()
        
        # Apply filters
        if package_type:
            queryset = queryset.filter(package__name__icontains=package_type)
        
        if status_filter:
            queryset = queryset.filter(status=status_filter)
        
        if search_query:
            queryset = queryset.filter(
                Q(patient__user__first_name__icontains=search_query) |
                Q(patient__user__last_name__icontains=search_query) |
                Q(patient__user__phone_number__icontains=search_query)
            )
        
        # Order by creation date (newest first)
        queryset = queryset.order_by('-created_at')
        
        # Paginate
        paginator = Paginator(queryset, page_size)
        page_obj = paginator.get_page(page)
        
        # Serialize data
        serializer = PatientSubscriptionSerializer(page_obj.object_list, many=True)
        
        # Prepare response data
        response_data = {
            'count': paginator.count,
            'total_pages': paginator.num_pages,
            'current_page': page,
            'page_size': page_size,
            'has_next': page_obj.has_next(),
            'has_previous': page_obj.has_previous(),
            'results': []
        }
        
        # Format results to match the design
        for subscription in serializer.data:
            # Get patient phone from user model
            patient_phone = None
            if subscription.get('patient'):
                try:
                    patient = Patient.objects.select_related('user').get(id=subscription['patient'])
                    patient_phone = patient.user.phone_number
                except Patient.DoesNotExist:
                    pass
            
            # Determine status display
            status_display = subscription.get('status', '').title()
            if subscription.get('is_active'):
                if subscription.get('end_date'):
                    from datetime import datetime, date
                    end_date = datetime.strptime(subscription['end_date'], '%Y-%m-%d').date()
                    days_until_expiry = (end_date - date.today()).days
                    if days_until_expiry <= 7:
                        status_display = "Expiring Soon"
                    else:
                        status_display = "Active"
            
            result_item = {
                'id': subscription['id'],
                'patient_name': subscription.get('patient_name', 'N/A'),
                'phone': patient_phone or 'N/A',
                'package': subscription.get('package_name', 'N/A'),
                'start_date': subscription.get('start_date'),
                'end_date': subscription.get('end_date'),
                'status': status_display,
                'consultations_used': subscription.get('consultations_used', 0),
                'consultations_remaining': subscription.get('consultations_remaining', 0),
            }
            response_data['results'].append(result_item)
        
        return Response(response_data)
    
    @swagger_auto_schema(
        operation_description="Send reminder notification to patient about subscription",
        request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            properties={
                'message': openapi.Schema(type=openapi.TYPE_STRING, description="Custom reminder message")
            }
        )
    )
    @action(detail=True, methods=['post'], url_path='remind')
    def send_reminder(self, request, pk=None):
        """
        Send reminder notification to patient about their subscription
        """
        try:
            subscription = PatientSubscription.objects.select_related(
                'patient__user', 'package'
            ).get(id=pk)
        except PatientSubscription.DoesNotExist:
            return Response({
                'error': 'Subscription not found'
            }, status=status.HTTP_404_NOT_FOUND)
        
        # Get custom message or use default
        custom_message = request.data.get('message', '')
        
        # Prepare notification data
        patient_name = subscription.patient.user.get_full_name() or subscription.patient.user.username
        package_name = subscription.package.name
        end_date = subscription.end_date.strftime('%B %d, %Y')
        
        if custom_message:
            message = custom_message
        else:
            message = f"Hi {patient_name}, your {package_name} subscription expires on {end_date}. Please renew to continue enjoying our services."
        
        # Here you would integrate with your notification service (email, SMS, push notifications)
        # For now, we'll just return a success response
        
        return Response({
            'success': True,
            'message': 'Reminder sent successfully',
            'details': {
                'patient_name': patient_name,
                'package_name': package_name,
                'end_date': end_date,
                'notification_message': message
            }
        })
    
    @swagger_auto_schema(
        operation_description="Export subscription data to CSV",
        manual_parameters=[
            openapi.Parameter('package_type', openapi.IN_QUERY, type=openapi.TYPE_STRING),
            openapi.Parameter('status', openapi.IN_QUERY, type=openapi.TYPE_STRING),
            openapi.Parameter('search', openapi.IN_QUERY, type=openapi.TYPE_STRING),
        ]
    )
    @action(detail=False, methods=['get'], url_path='export/csv')
    def export_csv(self, request):
        """
        Export subscription data to CSV file
        """
        import csv
        from django.http import HttpResponse
        from io import StringIO
        
        # Get filtered queryset using same logic as subscriptions endpoint
        package_type = request.query_params.get('package_type')
        status_filter = request.query_params.get('status')
        search_query = request.query_params.get('search')
        
        queryset = PatientSubscription.objects.select_related(
            'patient__user', 'package', 'payment'
        ).all()
        
        if package_type:
            queryset = queryset.filter(package__name__icontains=package_type)
        if status_filter:
            queryset = queryset.filter(status=status_filter)
        if search_query:
            queryset = queryset.filter(
                Q(patient__user__first_name__icontains=search_query) |
                Q(patient__user__last_name__icontains=search_query) |
                Q(patient__user__phone_number__icontains=search_query)
            )
        
        queryset = queryset.order_by('-created_at')
        
        # Create CSV response
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="package_subscriptions.csv"'
        
        writer = csv.writer(response)
        
        # Write headers
        writer.writerow([
            'Patient Name', 'Phone', 'Package', 'Start Date', 'End Date',
            'Status', 'Consultations Used', 'Consultations Remaining', 'Amount Paid'
        ])
        
        # Write data rows
        for subscription in queryset:
            patient_name = subscription.patient.user.get_full_name() or subscription.patient.user.username
            patient_phone = subscription.patient.user.phone_number or 'N/A'
            amount_paid = str(subscription.payment.amount) if subscription.payment else 'N/A'
            
            writer.writerow([
                patient_name,
                patient_phone,
                subscription.package.name,
                subscription.start_date,
                subscription.end_date,
                subscription.status.title(),
                subscription.consultations_used,
                subscription.consultations_remaining,
                amount_paid
            ])
        
        return response
    
    @swagger_auto_schema(
        operation_description="Export subscription data to PDF",
        manual_parameters=[
            openapi.Parameter('package_type', openapi.IN_QUERY, type=openapi.TYPE_STRING),
            openapi.Parameter('status', openapi.IN_QUERY, type=openapi.TYPE_STRING),
            openapi.Parameter('search', openapi.IN_QUERY, type=openapi.TYPE_STRING),
        ]
    )
    @action(detail=False, methods=['get'], url_path='export/pdf')
    def export_pdf(self, request):
        """
        Export subscription data to PDF file
        """
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from io import BytesIO
        from datetime import date
        
        # Get filtered queryset
        package_type = request.query_params.get('package_type')
        status_filter = request.query_params.get('status')
        search_query = request.query_params.get('search')
        
        queryset = PatientSubscription.objects.select_related(
            'patient__user', 'package', 'payment'
        ).all()
        
        if package_type:
            queryset = queryset.filter(package__name__icontains=package_type)
        if status_filter:
            queryset = queryset.filter(status=status_filter)
        if search_query:
            queryset = queryset.filter(
                Q(patient__user__first_name__icontains=search_query) |
                Q(patient__user__last_name__icontains=search_query) |
                Q(patient__user__phone_number__icontains=search_query)
            )
        
        queryset = queryset.order_by('-created_at')
        
        # Create PDF
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        elements = []
        
        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1,  # Center alignment
        )
        
        # Title
        title = Paragraph("Package Payment Tracker Report", title_style)
        elements.append(title)
        elements.append(Spacer(1, 12))
        
        # Report date
        date_para = Paragraph(f"Generated on: {date.today().strftime('%B %d, %Y')}", styles['Normal'])
        elements.append(date_para)
        elements.append(Spacer(1, 12))
        
        # Table data
        data = [['Patient Name', 'Package', 'Status', 'Start Date', 'End Date', 'Consultations']]
        
        for subscription in queryset:
            patient_name = subscription.patient.user.get_full_name() or subscription.patient.user.username
            data.append([
                patient_name,
                subscription.package.name,
                subscription.status.title(),
                subscription.start_date.strftime('%Y-%m-%d'),
                subscription.end_date.strftime('%Y-%m-%d'),
                f"{subscription.consultations_used}/{subscription.package.consultation_limit}"
            ])
        
        # Create table
        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        elements.append(table)
        
        doc.build(elements)
        buffer.seek(0)
        
        response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="package_subscriptions.pdf"'
        
        return response


class RiskSegmentationViewSet(viewsets.ViewSet):
    """
    ViewSet for risk segmentation analytics and reporting
    """
    permission_classes = [permissions.IsAuthenticated]
    
    @swagger_auto_schema(
        operation_description="Get risk segmentation summary with distribution statistics",
        manual_parameters=[
            openapi.Parameter(
                'county', openapi.IN_QUERY, description="Filter by county/city", 
                type=openapi.TYPE_STRING, required=False
            ),
            openapi.Parameter(
                'date_from', openapi.IN_QUERY, description="Filter from date (YYYY-MM-DD)", 
                type=openapi.TYPE_STRING, format=openapi.FORMAT_DATE, required=False
            ),
            openapi.Parameter(
                'date_to', openapi.IN_QUERY, description="Filter to date (YYYY-MM-DD)", 
                type=openapi.TYPE_STRING, format=openapi.FORMAT_DATE, required=False
            )
        ],
        responses={200: RiskSegmentationSummarySerializer}
    )
    @action(detail=False, methods=['get'])
    def summary(self, request):
        """
        Get risk segmentation summary statistics
        """
        from django.db.models import Count, Case, When, F, Max
        from django.db import models
        from datetime import datetime
        
        # Base queryset - get latest appointments per patient
        appointments = Appointment.objects.filter(
            risk_level__isnull=False
        ).exclude(risk_level='').select_related(
            'patient__user', 'healthcare_facility'
        )
        
        # Apply filters
        county = request.query_params.get('county')
        date_from = request.query_params.get('date_from')
        date_to = request.query_params.get('date_to')
        
        date_range = {}
        
        if county:
            appointments = appointments.filter(
                Q(healthcare_facility__city__icontains=county) |
                Q(patient__user__address__icontains=county)
            )
        
        if date_from:
            try:
                date_from_parsed = datetime.strptime(date_from, '%Y-%m-%d').date()
                appointments = appointments.filter(appointment_date__gte=date_from_parsed)
                date_range['from'] = date_from
            except ValueError:
                pass
        
        if date_to:
            try:
                date_to_parsed = datetime.strptime(date_to, '%Y-%m-%d').date()
                appointments = appointments.filter(appointment_date__lte=date_to_parsed)
                date_range['to'] = date_to
            except ValueError:
                pass
        
        # Get latest appointment per patient to avoid duplicates
        latest_appointments = appointments.values('patient_id').annotate(
            latest_date=models.Max('appointment_date')
        )
        
        # Filter appointments to only include the latest per patient
        filtered_appointments = appointments.filter(
            models.Q(patient_id__in=[item['patient_id'] for item in latest_appointments])
        )
        
        # For each patient, get their most recent appointment
        patient_risk_data = {}
        for appointment in filtered_appointments.order_by('patient_id', '-appointment_date'):
            if appointment.patient_id not in patient_risk_data:
                patient_risk_data[appointment.patient_id] = appointment.risk_level
        
        # Count risk levels
        risk_counts = {}
        total_patients = len(patient_risk_data)
        
        for risk_level in patient_risk_data.values():
            # Normalize risk levels to match design (severe, high, moderate)
            normalized_risk = self._normalize_risk_level(risk_level)
            risk_counts[normalized_risk] = risk_counts.get(normalized_risk, 0) + 1
        
        # Calculate percentages and create distribution
        distribution = []
        for risk_level in ['severe', 'high', 'moderate']:
            count = risk_counts.get(risk_level, 0)
            percentage = round((count / total_patients * 100), 1) if total_patients > 0 else 0
            
            distribution.append({
                'risk_level': risk_level,
                'patient_count': count,
                'percentage': percentage
            })
        
        # Prepare response data
        response_data = {
            'total_patients': total_patients,
            'distribution': distribution
        }
        
        if date_range:
            response_data['date_range'] = date_range
        if county:
            response_data['county_filter'] = county
        
        serializer = RiskSegmentationSummarySerializer(data=response_data)
        serializer.is_valid()
        return Response(serializer.data)
    
    def _normalize_risk_level(self, risk_level):
        """
        Normalize risk levels to match the design requirements
        """
        if not risk_level:
            return 'moderate'
        
        risk_level = risk_level.lower()
        if risk_level in ['critical', 'severe']:
            return 'severe'
        elif risk_level in ['high']:
            return 'high' 
        elif risk_level in ['medium', 'moderate', 'low']:
            return 'moderate'
        else:
            return 'moderate'
    
    @swagger_auto_schema(
        operation_description="Get list of patients by risk level",
        manual_parameters=[
            openapi.Parameter(
                'risk_level', openapi.IN_QUERY, 
                description="Filter by risk level (severe, high, moderate)", 
                type=openapi.TYPE_STRING, required=True,
                enum=['severe', 'high', 'moderate']
            ),
            openapi.Parameter(
                'county', openapi.IN_QUERY, description="Filter by county/city", 
                type=openapi.TYPE_STRING, required=False
            ),
            openapi.Parameter(
                'date_from', openapi.IN_QUERY, description="Filter from date (YYYY-MM-DD)", 
                type=openapi.TYPE_STRING, format=openapi.FORMAT_DATE, required=False
            ),
            openapi.Parameter(
                'date_to', openapi.IN_QUERY, description="Filter to date (YYYY-MM-DD)", 
                type=openapi.TYPE_STRING, format=openapi.FORMAT_DATE, required=False
            )
        ],
        responses={200: PatientRiskListSerializer(many=True)}
    )
    @action(detail=False, methods=['get'])
    def patients(self, request):
        """
        Get list of patients filtered by risk level
        """
        from datetime import datetime
        
        risk_level = request.query_params.get('risk_level')
        if not risk_level:
            return Response({
                'error': 'risk_level parameter is required'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # Base queryset
        appointments = Appointment.objects.filter(
            risk_level__isnull=False
        ).exclude(risk_level='').select_related(
            'patient__user', 'doctor__user', 'healthcare_facility'
        )
        
        # Apply filters
        county = request.query_params.get('county')
        date_from = request.query_params.get('date_from')
        date_to = request.query_params.get('date_to')
        
        if county:
            appointments = appointments.filter(
                Q(healthcare_facility__city__icontains=county) |
                Q(patient__user__address__icontains=county)
            )
        
        if date_from:
            try:
                date_from_parsed = datetime.strptime(date_from, '%Y-%m-%d').date()
                appointments = appointments.filter(appointment_date__gte=date_from_parsed)
            except ValueError:
                pass
        
        if date_to:
            try:
                date_to_parsed = datetime.strptime(date_to, '%Y-%m-%d').date()
                appointments = appointments.filter(appointment_date__lte=date_to_parsed)
            except ValueError:
                pass
        
        # Filter by risk level - handle both original and normalized values
        risk_level_filters = self._get_risk_level_filters(risk_level)
        appointments = appointments.filter(risk_level__in=risk_level_filters)
        
        # Get latest appointment per patient
        patient_latest_appointments = {}
        for appointment in appointments.order_by('patient_id', '-appointment_date'):
            if appointment.patient_id not in patient_latest_appointments:
                patient_latest_appointments[appointment.patient_id] = appointment
        
        # Convert to list and apply pagination
        latest_appointments = list(patient_latest_appointments.values())
        
        # Serialize and return
        serializer = PatientRiskListSerializer(latest_appointments, many=True)
        return Response(serializer.data)
    
    def _get_risk_level_filters(self, normalized_risk_level):
        """
        Get list of original risk levels that map to the normalized level
        """
        mapping = {
            'severe': ['critical', 'severe'],
            'high': ['high'],
            'moderate': ['medium', 'moderate', 'low']
        }
        return mapping.get(normalized_risk_level.lower(), [])
    
    @swagger_auto_schema(
        operation_description="Export risk segmentation data to PDF",
        manual_parameters=[
            openapi.Parameter(
                'county', openapi.IN_QUERY, description="Filter by county/city", 
                type=openapi.TYPE_STRING, required=False
            ),
            openapi.Parameter(
                'date_from', openapi.IN_QUERY, description="Filter from date (YYYY-MM-DD)", 
                type=openapi.TYPE_STRING, format=openapi.FORMAT_DATE, required=False
            ),
            openapi.Parameter(
                'date_to', openapi.IN_QUERY, description="Filter to date (YYYY-MM-DD)", 
                type=openapi.TYPE_STRING, format=openapi.FORMAT_DATE, required=False
            )
        ]
    )
    @action(detail=False, methods=['get'])
    def export_pdf(self, request):
        """
        Export risk segmentation report to PDF
        """
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        from datetime import datetime
        import io
        
        # Get the summary data
        summary_response = self.summary(request)
        summary_data = summary_response.data
        
        # Create PDF buffer
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        elements = []
        styles = getSampleStyleSheet()
        
        # Title
        title = Paragraph("Risk Segmentation Report", styles['Title'])
        elements.append(title)
        elements.append(Spacer(1, 12))
        
        # Filters applied
        filters_text = "Filters Applied: "
        if summary_data.get('county_filter'):
            filters_text += f"County: {summary_data['county_filter']}, "
        if summary_data.get('date_range'):
            date_range = summary_data['date_range']
            if date_range.get('from'):
                filters_text += f"From: {date_range['from']}, "
            if date_range.get('to'):
                filters_text += f"To: {date_range['to']}, "
        if filters_text.endswith(", "):
            filters_text = filters_text[:-2]
        
        elements.append(Paragraph(filters_text, styles['Normal']))
        elements.append(Spacer(1, 12))
        
        # Summary statistics
        summary_text = f"Total Patients: {summary_data['total_patients']}"
        elements.append(Paragraph(summary_text, styles['Heading2']))
        elements.append(Spacer(1, 12))
        
        # Distribution table
        data = [["Risk Level", "Patient Count", "Percentage"]]
        for item in summary_data['distribution']:
            data.append([
                item['risk_level'].title(),
                str(item['patient_count']),
                f"{item['percentage']}%"
            ])
        
        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        elements.append(table)
        
        doc.build(elements)
        buffer.seek(0)
        
        response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="risk_segmentation_report.pdf"'
        
        return response
    
    @swagger_auto_schema(
        operation_description="Export risk segmentation data to CSV",
        manual_parameters=[
            openapi.Parameter(
                'county', openapi.IN_QUERY, description="Filter by county/city", 
                type=openapi.TYPE_STRING, required=False
            ),
            openapi.Parameter(
                'date_from', openapi.IN_QUERY, description="Filter from date (YYYY-MM-DD)", 
                type=openapi.TYPE_STRING, format=openapi.FORMAT_DATE, required=False
            ),
            openapi.Parameter(
                'date_to', openapi.IN_QUERY, description="Filter to date (YYYY-MM-DD)", 
                type=openapi.TYPE_STRING, format=openapi.FORMAT_DATE, required=False
            )
        ]
    )
    @action(detail=False, methods=['get'])
    def export_csv(self, request):
        """
        Export risk segmentation detailed data to CSV
        """
        import csv
        from django.http import HttpResponse
        from datetime import datetime
        
        # Get all patients data for each risk level
        all_patients = []
        for risk_level in ['severe', 'high', 'moderate']:
            # Create a new request with the risk_level parameter
            modified_params = request.query_params.copy()
            modified_params['risk_level'] = risk_level
            
            # Temporarily modify the request to get patients for each risk level
            request.query_params = modified_params
            patients_response = self.patients(request)
            
            if patients_response.status_code == 200:
                for patient_data in patients_response.data:
                    all_patients.append(patient_data)
        
        # Create CSV response
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="risk_segmentation_patients.csv"'
        
        writer = csv.writer(response)
        
        # Write header
        writer.writerow([
            'Patient Name', 'Email', 'Age', 'County', 'Risk Level', 
            'Latest Appointment Date', 'Doctor', 'Diagnosis', 'Treatment'
        ])
        
        # Write data rows
        for patient in all_patients:
            writer.writerow([
                patient.get('patient_name', ''),
                patient.get('patient_email', ''),
                patient.get('age', ''),
                patient.get('county', ''),
                patient.get('risk_level', '').title(),
                patient.get('latest_appointment_date', ''),
                patient.get('doctor_name', ''),
                patient.get('diagnosis', ''),
                patient.get('treatment', '')
            ])
        
        return response


class TeleconsultationLogViewSet(viewsets.ModelViewSet):
    """
    ViewSet for teleconsultation logs with statistics and filtering
    """
    queryset = Consultation.objects.all()
    serializer_class = ConsultationSerializer
    permission_classes = [permissions.IsAuthenticated]
    pagination_class = CustomPageNumberPagination
    
    def get_permissions(self):
        if self.action in ['destroy']:
            permission_classes = [IsAdminUser]
        else:
            permission_classes = [permissions.IsAuthenticated]
        return [permission() for permission in permission_classes]
    
    def get_queryset(self):
        queryset = Consultation.objects.select_related(
            'appointment__patient__user', 
            'appointment__doctor__user', 
            'appointment__healthcare_facility'
        ).all()
        
        # Role-based filtering
        if self.request.user.is_authenticated and self.request.user.roles.filter(name='admin').exists():
            pass
        elif self.request.user.is_authenticated and self.request.user.roles.filter(name='doctor').exists():
            try:
                doctor = self.request.user.doctor
                queryset = queryset.filter(appointment__doctor=doctor)
            except Doctor.DoesNotExist:
                return Consultation.objects.none()
        elif self.request.user.is_authenticated and self.request.user.roles.filter(name='patient').exists():
            try:
                patient = self.request.user.patient
                queryset = queryset.filter(appointment__patient=patient)
            except Patient.DoesNotExist:
                return Consultation.objects.none()
        else:
            return Consultation.objects.none()
            
        # Apply filters
        patient_name = self.request.query_params.get('patient_name')
        if patient_name:
            queryset = queryset.filter(
                Q(appointment__patient__user__first_name__icontains=patient_name) |
                Q(appointment__patient__user__last_name__icontains=patient_name)
            )
            
        clinic_name = self.request.query_params.get('clinic')
        if clinic_name:
            queryset = queryset.filter(appointment__healthcare__name__icontains=clinic_name)
            
        consultation_type = self.request.query_params.get('consultation_type')
        if consultation_type:
            if consultation_type.lower() == 'video':
                queryset = queryset.filter(
                    appointment__appointment_type__in=['video', 'teleconsultation']
                )
            elif consultation_type.lower() == 'missed':
                queryset = queryset.filter(status='missed')
                
        date_from = self.request.query_params.get('date_from')
        if date_from:
            try:
                date_from = datetime.strptime(date_from, '%Y-%m-%d').date()
                queryset = queryset.filter(appointment__appointment_date__gte=date_from)
            except ValueError:
                pass
                
        date_to = self.request.query_params.get('date_to')
        if date_to:
            try:
                date_to = datetime.strptime(date_to, '%Y-%m-%d').date()
                queryset = queryset.filter(appointment__appointment_date__lte=date_to)
            except ValueError:
                pass
                
        return queryset.order_by('-appointment__appointment_date', '-start_time')
    
    def list(self, request, *args, **kwargs):
        queryset = self.get_queryset()
        
        # Check if summary is requested
        include_summary = request.query_params.get('summary', 'false').lower() == 'true'
        
        if include_summary:
            # Calculate summary statistics
            total_consultations = queryset.count()
            completed_consultations = queryset.filter(status='completed').count()
            missed_consultations = queryset.filter(status='missed').count()
            upcoming_consultations = queryset.filter(status='scheduled').count()
            in_progress_consultations = queryset.filter(status='in-progress').count()
            cancelled_consultations = queryset.filter(status='cancelled').count()
            
            # Calculate rates
            completion_rate = (completed_consultations / total_consultations * 100) if total_consultations > 0 else 0
            miss_rate = (missed_consultations / total_consultations * 100) if total_consultations > 0 else 0
            
            summary = {
                'total_consultations': total_consultations,
                'completed_consultations': completed_consultations,
                'missed_consultations': missed_consultations,
                'upcoming_consultations': upcoming_consultations,
                'in_progress_consultations': in_progress_consultations,
                'cancelled_consultations': cancelled_consultations,
                'completion_rate': round(completion_rate, 1),
                'miss_rate': round(miss_rate, 1)
            }
            
            # Get paginated results
            page = self.paginate_queryset(queryset)
            if page is not None:
                serializer = self.get_serializer(page, many=True)
                paginated_response = self.get_paginated_response(serializer.data)
                paginated_response.data['summary'] = summary
                return paginated_response
            
            # If no pagination
            serializer = self.get_serializer(queryset, many=True)
            return Response({
                'results': serializer.data,
                'summary': summary
            })
        
        # Default behavior without summary
        return super().list(request, *args, **kwargs)
    
    @swagger_auto_schema(
        operation_description="Get teleconsultation statistics",
        responses={
            200: openapi.Response("Statistics", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'total_consultations': openapi.Schema(type=openapi.TYPE_INTEGER),
                    'completed_consultations': openapi.Schema(type=openapi.TYPE_INTEGER),
                    'missed_consultations': openapi.Schema(type=openapi.TYPE_INTEGER),
                    'upcoming_consultations': openapi.Schema(type=openapi.TYPE_INTEGER),
                    'completion_rate': openapi.Schema(type=openapi.TYPE_NUMBER),
                    'miss_rate': openapi.Schema(type=openapi.TYPE_NUMBER)
                }
            ))
        }
    )
    @action(detail=False, methods=['get'])
    def statistics(self, request):
        """
        Get teleconsultation statistics
        """
        queryset = self.get_queryset()
        
        total = queryset.count()
        completed = queryset.filter(status='completed').count()
        missed = queryset.filter(status='missed').count()
        upcoming = queryset.filter(
            status__in=['scheduled', 'in-progress'],
            appointment__appointment_date__gte=timezone.now().date()
        ).count()
        
        completion_rate = (completed / total * 100) if total > 0 else 0
        miss_rate = (missed / total * 100) if total > 0 else 0
        
        return Response({
            'total_consultations': total,
            'completed_consultations': completed,
            'missed_consultations': missed,
            'upcoming_consultations': upcoming,
            'completion_rate': round(completion_rate, 1),
            'miss_rate': round(miss_rate, 1)
        })
    
    @swagger_auto_schema(
        operation_description="Export teleconsultation logs to PDF",
        manual_parameters=[
            openapi.Parameter('patient_name', openapi.IN_QUERY, type=openapi.TYPE_STRING),
            openapi.Parameter('clinic', openapi.IN_QUERY, type=openapi.TYPE_STRING),
            openapi.Parameter('consultation_type', openapi.IN_QUERY, type=openapi.TYPE_STRING),
            openapi.Parameter('date_from', openapi.IN_QUERY, type=openapi.TYPE_STRING, format=openapi.FORMAT_DATE),
            openapi.Parameter('date_to', openapi.IN_QUERY, type=openapi.TYPE_STRING, format=openapi.FORMAT_DATE)
        ]
    )
    @action(detail=False, methods=['get'])
    def export_pdf(self, request):
        """
        Export teleconsultation logs to PDF
        """
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib import colors
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER
        import io
        
        queryset = self.get_queryset()
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        elements = []
        
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        elements.append(Paragraph("Teleconsultation Log Report", title_style))
        elements.append(Spacer(1, 20))
        
        # Table data
        data = [['Date', 'Time', 'Patient', 'Doctor', 'Clinic', 'Type', 'Status']]
        
        for consultation in queryset:
            appointment = consultation.appointment
            data.append([
                str(appointment.appointment_date),
                str(appointment.start_time) if appointment.start_time else 'N/A',
                f"{appointment.patient.user.first_name} {appointment.patient.user.last_name}",
                f"Dr. {appointment.doctor.user.first_name} {appointment.doctor.user.last_name}",
                appointment.healthcare.name,
                appointment.appointment_type.title(),
                consultation.status.title()
            ])
        
        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        elements.append(table)
        doc.build(elements)
        buffer.seek(0)
        
        response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="teleconsultation_logs.pdf"'
        
        return response
    
    @swagger_auto_schema(
        operation_description="Export teleconsultation logs to CSV"
    )
    @action(detail=False, methods=['get'])
    def export_csv(self, request):
        """
        Export teleconsultation logs to CSV
        """
        import csv
        
        queryset = self.get_queryset()
        
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="teleconsultation_logs.csv"'
        
        writer = csv.writer(response)
        writer.writerow(['Date', 'Time', 'Patient', 'Doctor', 'Clinic', 'Type', 'Status', 'Duration'])
        
        for consultation in queryset:
            appointment = consultation.appointment
            duration = ""
            if consultation.start_time and consultation.end_time:
                duration_delta = consultation.end_time - consultation.start_time
                duration = str(duration_delta)
                
            writer.writerow([
                appointment.appointment_date,
                appointment.start_time or 'N/A',
                f"{appointment.patient.user.first_name} {appointment.patient.user.last_name}",
                f"Dr. {appointment.doctor.user.first_name} {appointment.doctor.user.last_name}",
                appointment.healthcare.name,
                appointment.appointment_type.title(),
                consultation.status.title(),
                duration
            ])
        
        return response


class FollowUpComplianceViewSet(viewsets.ReadOnlyModelViewSet):
    """
    ViewSet for follow-up compliance monitoring
    """
    queryset = Appointment.objects.all()
    serializer_class = AppointmentSerializer
    permission_classes = [permissions.IsAuthenticated]
    pagination_class = CustomPageNumberPagination
    
    def get_queryset(self):
        queryset = Appointment.objects.select_related(
            'patient__user', 
            'doctor__user', 
            'healthcare_facility'
        ).all()
        
        # Role-based filtering
        if self.request.user.is_authenticated and self.request.user.roles.filter(name='admin').exists():
            pass
        elif self.request.user.is_authenticated and self.request.user.roles.filter(name='doctor').exists():
            try:
                doctor = self.request.user.doctor
                queryset = queryset.filter(doctor=doctor)
            except Doctor.DoesNotExist:
                return Appointment.objects.none()
        elif self.request.user.is_authenticated and self.request.user.roles.filter(name='patient').exists():
            try:
                patient = self.request.user.patient
                queryset = queryset.filter(patient=patient)
            except Patient.DoesNotExist:
                return Appointment.objects.none()
        else:
            return Appointment.objects.none()
        
        # Apply filters
        patient_name = self.request.query_params.get('patient_name')
        if patient_name:
            queryset = queryset.filter(
                Q(patient__user__first_name__icontains=patient_name) |
                Q(patient__user__last_name__icontains=patient_name)
            )
            
        risk_level = self.request.query_params.get('risk_level')
        if risk_level:
            queryset = queryset.filter(risk_level=risk_level)
            
        follow_up_type = self.request.query_params.get('follow_up_type')
        if follow_up_type:
            if follow_up_type.lower() == 'appointment':
                queryset = queryset.filter(appointment_type__in=['follow-up', 'routine'])
            elif follow_up_type.lower() == 'medication':
                queryset = queryset.filter(notes__icontains='medication')
                
        compliance_status = self.request.query_params.get('status')
        if compliance_status:
            if compliance_status.lower() == 'missed':
                queryset = queryset.filter(status='noshow')
                
        return queryset.order_by('-appointment_date', '-start_time')
    
    def list(self, request, *args, **kwargs):
        queryset = self.get_queryset()
        
        # Check if summary is requested
        include_summary = request.query_params.get('summary', 'false').lower() == 'true'
        
        if include_summary:
            # Calculate summary statistics
            total_appointments = queryset.count()
            fulfilled_appointments = queryset.filter(status='fulfilled').count()
            missed_appointments = queryset.filter(status='noshow').count()
            pending_appointments = queryset.filter(status='scheduled').count()
            
            compliance_rate = (fulfilled_appointments / total_appointments * 100) if total_appointments > 0 else 0
            
            # Critical non-compliant patients
            critical_non_compliant = queryset.filter(
                risk_level__in=['high', 'severe'],
                status='noshow'
            ).values('patient').distinct().count()
            
            summary = {
                'total_appointments': total_appointments,
                'fulfilled_appointments': fulfilled_appointments,
                'missed_appointments': missed_appointments,
                'pending_appointments': pending_appointments,
                'compliance_rate': round(compliance_rate, 1),
                'critical_non_compliant_patients': critical_non_compliant
            }
            
            # Get paginated results
            page = self.paginate_queryset(queryset)
            if page is not None:
                serializer = self.get_serializer(page, many=True)
                paginated_response = self.get_paginated_response(serializer.data)
                paginated_response.data['summary'] = summary
                return paginated_response
            
            # If no pagination
            serializer = self.get_serializer(queryset, many=True)
            return Response({
                'results': serializer.data,
                'summary': summary
            })
        
        # Default behavior without summary
        return super().list(request, *args, **kwargs)
    
    @swagger_auto_schema(
        operation_description="Get follow-up compliance statistics",
        responses={
            200: openapi.Response("Statistics", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'total_patients_tracked': openapi.Schema(type=openapi.TYPE_INTEGER),
                    'missed_appointments': openapi.Schema(type=openapi.TYPE_INTEGER),
                    'missed_med_reminders': openapi.Schema(type=openapi.TYPE_INTEGER),
                    'critical_non_compliant': openapi.Schema(type=openapi.TYPE_INTEGER),
                    'compliance_rate': openapi.Schema(type=openapi.TYPE_NUMBER)
                }
            ))
        }
    )
    @action(detail=False, methods=['get'])
    def statistics(self, request):
        """
        Get follow-up compliance statistics
        """
        queryset = self.get_queryset()
        
        total_patients = queryset.values('patient').distinct().count()
        missed_appointments = queryset.filter(status='noshow').count()
        
        # Count medication-related missed follow-ups (based on notes containing medication keywords)
        missed_med_reminders = queryset.filter(
            Q(notes__icontains='medication') | Q(notes__icontains='prescription'),
            status='noshow'
        ).count()
        
        # Critical non-compliant (high/severe risk patients with missed appointments)
        critical_non_compliant = queryset.filter(
            risk_level__in=['high', 'severe'],
            status='noshow'
        ).values('patient').distinct().count()
        
        # Calculate compliance rate
        total_appointments = queryset.count()
        fulfilled_appointments = queryset.filter(status='fulfilled').count()
        compliance_rate = (fulfilled_appointments / total_appointments * 100) if total_appointments > 0 else 0
        
        return Response({
            'total_patients_tracked': total_patients,
            'missed_appointments': missed_appointments,
            'missed_med_reminders': missed_med_reminders,
            'critical_non_compliant': critical_non_compliant,
            'compliance_rate': round(compliance_rate, 1)
        })
    
    @swagger_auto_schema(
        operation_description="Get list of patients for compliance monitoring",
        manual_parameters=[
            openapi.Parameter('patient_name', openapi.IN_QUERY, type=openapi.TYPE_STRING),
            openapi.Parameter('risk_level', openapi.IN_QUERY, type=openapi.TYPE_STRING),
            openapi.Parameter('follow_up_type', openapi.IN_QUERY, type=openapi.TYPE_STRING),
            openapi.Parameter('status', openapi.IN_QUERY, type=openapi.TYPE_STRING)
        ]
    )
    @action(detail=False, methods=['get'])
    def patients(self, request):
        """
        Get detailed patient compliance data
        """
        queryset = self.get_queryset()
        
        # Group by patient and calculate compliance metrics
        patient_data = []
        patient_appointments = {}
        
        for appointment in queryset:
            patient_id = appointment.patient.id
            if patient_id not in patient_appointments:
                patient_appointments[patient_id] = {
                    'patient': appointment.patient,
                    'appointments': [],
                    'missed_count': 0,
                    'total_count': 0
                }
            
            patient_appointments[patient_id]['appointments'].append(appointment)
            patient_appointments[patient_id]['total_count'] += 1
            
            if appointment.status == 'noshow':
                patient_appointments[patient_id]['missed_count'] += 1
        
        for patient_id, data in patient_appointments.items():
            patient = data['patient']
            latest_appointment = max(data['appointments'], key=lambda x: x.appointment_date)
            
            patient_data.append({
                'patient_name': f"{patient.user.first_name} {patient.user.last_name}",
                'phone': patient.user.email,  # Using email as phone not available
                'risk_level': latest_appointment.risk_level,
                'missed_count': data['missed_count'],
                'follow_up_type': 'Appointment' if latest_appointment.appointment_type in ['follow-up', 'routine'] else 'Medication',
                'status': 'Missed' if data['missed_count'] > 0 else 'Compliant'
            })
        
        return Response(patient_data)
    
    @swagger_auto_schema(
        operation_description="Export compliance data to PDF"
    )
    @action(detail=False, methods=['get'])
    def export_pdf(self, request):
        """
        Export follow-up compliance data to PDF
        """
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER
        import io
        
        patients_response = self.patients(request)
        patients_data = patients_response.data
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        elements = []
        
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        elements.append(Paragraph("Follow-Up Compliance Report", title_style))
        elements.append(Spacer(1, 20))
        
        data = [['Patient Name', 'Phone', 'Risk Level', 'Missed Count', 'Follow-Up Type', 'Status']]
        
        for patient in patients_data:
            data.append([
                patient['patient_name'],
                patient['phone'],
                patient['risk_level'].title(),
                str(patient['missed_count']),
                patient['follow_up_type'],
                patient['status']
            ])
        
        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        elements.append(table)
        doc.build(elements)
        buffer.seek(0)
        
        response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="compliance_report.pdf"'
        
        return response
    
    @swagger_auto_schema(
        operation_description="Export compliance data to CSV"
    )
    @action(detail=False, methods=['get'])
    def export_csv(self, request):
        """
        Export follow-up compliance data to CSV
        """
        import csv
        
        patients_response = self.patients(request)
        patients_data = patients_response.data
        
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="compliance_data.csv"'
        
        writer = csv.writer(response)
        writer.writerow(['Patient Name', 'Phone', 'Risk Level', 'Missed Count', 'Follow-Up Type', 'Status'])
        
        for patient in patients_data:
            writer.writerow([
                patient['patient_name'],
                patient['phone'],
                patient['risk_level'].title(),
                patient['missed_count'],
                patient['follow_up_type'],
                patient['status']
            ])
        
        return response


class EnhancedAppointmentListViewSet(viewsets.ReadOnlyModelViewSet):
    """
    Enhanced appointment list with advanced search and filtering
    """
    queryset = Appointment.objects.all()
    serializer_class = AppointmentSerializer
    permission_classes = [permissions.IsAuthenticated]
    pagination_class = CustomPageNumberPagination
    
    def get_queryset(self):
        queryset = Appointment.objects.select_related(
            'patient__user', 
            'doctor__user', 
            'healthcare_facility'
        ).all()
        
        # Role-based filtering
        if self.request.user.is_authenticated and self.request.user.roles.filter(name='admin').exists():
            pass
        elif self.request.user.is_authenticated and self.request.user.roles.filter(name='doctor').exists():
            try:
                doctor = self.request.user.doctor
                queryset = queryset.filter(doctor=doctor)
            except Doctor.DoesNotExist:
                return Appointment.objects.none()
        elif self.request.user.is_authenticated and self.request.user.roles.filter(name='patient').exists():
            try:
                patient = self.request.user.patient
                queryset = queryset.filter(patient=patient)
            except Patient.DoesNotExist:
                return Appointment.objects.none()
        else:
            return Appointment.objects.none()
        
        # Apply search and filters
        search = self.request.query_params.get('search')
        if search:
            queryset = queryset.filter(
                Q(patient__user__first_name__icontains=search) |
                Q(patient__user__last_name__icontains=search) |
                Q(doctor__user__first_name__icontains=search) |
                Q(doctor__user__last_name__icontains=search) |
                Q(healthcare_facility__name__icontains=search)
            )
            
        status_filter = self.request.query_params.get('status')
        if status_filter:
            queryset = queryset.filter(status=status_filter)
            
        appointment_type = self.request.query_params.get('type')
        if appointment_type:
            queryset = queryset.filter(appointment_type=appointment_type)
            
        subscription_type = self.request.query_params.get('subscription')
        if subscription_type:
            # Filter by patient's subscription type
            queryset = queryset.filter(
                patient__patientsubscription__package__name__icontains=subscription_type,
                patient__patientsubscription__status='active'
            )
        
        date_from = self.request.query_params.get('date_from')
        if date_from:
            try:
                date_from = datetime.strptime(date_from, '%Y-%m-%d').date()
                queryset = queryset.filter(appointment_date__gte=date_from)
            except ValueError:
                pass
                
        date_to = self.request.query_params.get('date_to')
        if date_to:
            try:
                date_to = datetime.strptime(date_to, '%Y-%m-%d').date()
                queryset = queryset.filter(appointment_date__lte=date_to)
            except ValueError:
                pass
        
        return queryset.order_by('-appointment_date', '-start_time')
    
    @swagger_auto_schema(
        operation_description="Export appointment list to PDF",
        manual_parameters=[
            openapi.Parameter('search', openapi.IN_QUERY, type=openapi.TYPE_STRING),
            openapi.Parameter('status', openapi.IN_QUERY, type=openapi.TYPE_STRING),
            openapi.Parameter('type', openapi.IN_QUERY, type=openapi.TYPE_STRING),
            openapi.Parameter('subscription', openapi.IN_QUERY, type=openapi.TYPE_STRING)
        ]
    )
    @action(detail=False, methods=['get'])
    def export_pdf(self, request):
        """
        Export appointment list to PDF
        """
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER
        import io
        
        queryset = self.get_queryset()
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        elements = []
        
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        elements.append(Paragraph("Appointments List Report", title_style))
        elements.append(Spacer(1, 20))
        
        data = [['Name', 'Time', 'Type', 'Subscription', 'Status']]
        
        for appointment in queryset:
            # Get subscription info
            subscription = "N/A"
            try:
                active_subscription = appointment.patient.patientsubscription_set.filter(
                    status='active'
                ).first()
                if active_subscription:
                    subscription = active_subscription.package.name
            except:
                pass
                
            data.append([
                f"{appointment.patient.user.first_name} {appointment.patient.user.last_name}",
                f"{appointment.appointment_date} {appointment.start_time or 'N/A'}",
                appointment.appointment_type.title(),
                subscription,
                appointment.status.title()
            ])
        
        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        elements.append(table)
        doc.build(elements)
        buffer.seek(0)
        
        response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="appointments_list.pdf"'
        
        return response
    
    @swagger_auto_schema(
        operation_description="Export appointment list to CSV"
    )
    @action(detail=False, methods=['get'])
    def export_csv(self, request):
        """
        Export appointment list to CSV
        """
        import csv
        
        queryset = self.get_queryset()
        
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="appointments_list.csv"'
        
        writer = csv.writer(response)
        writer.writerow(['Name', 'Time', 'Type', 'Subscription', 'Status', 'Doctor', 'Clinic'])
        
        for appointment in queryset:
            # Get subscription info
            subscription = "N/A"
            try:
                active_subscription = appointment.patient.patientsubscription_set.filter(
                    status='active'
                ).first()
                if active_subscription:
                    subscription = active_subscription.package.name
            except:
                pass
                
            writer.writerow([
                f"{appointment.patient.user.first_name} {appointment.patient.user.last_name}",
                f"{appointment.appointment_date} {appointment.start_time or 'N/A'}",
                appointment.appointment_type.title(),
                subscription,
                appointment.status.title(),
                f"Dr. {appointment.doctor.user.first_name} {appointment.doctor.user.last_name}",
                appointment.healthcare_facility.name if appointment.healthcare_facility else 'N/A'
            ])
        
        return response


class PatientJournalViewSet(viewsets.ModelViewSet):
    """
    ViewSet for managing patient journals
    """
    queryset = PatientJournal.objects.all()
    serializer_class = PatientJournalSerializer
    permission_classes = [permissions.IsAuthenticated]
    pagination_class = CustomPageNumberPagination
    
    @swagger_auto_schema(
        operation_description="List all journals with pagination and meta data",
        responses={
            200: openapi.Response("Success", openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'data': openapi.Schema(
                        type=openapi.TYPE_ARRAY,
                        items=openapi.Schema(type=openapi.TYPE_OBJECT)
                    ),
                    'meta': openapi.Schema(
                        type=openapi.TYPE_OBJECT,
                        properties={
                            'total_count': openapi.Schema(type=openapi.TYPE_INTEGER),
                            'page': openapi.Schema(type=openapi.TYPE_INTEGER),
                            'per_page': openapi.Schema(type=openapi.TYPE_INTEGER),
                        }
                    )
                }
            ))
        }
    )
    def list(self, request, *args, **kwargs):
        """
        List journals with custom pagination format
        """
        queryset = self.filter_queryset(self.get_queryset())
        page = self.paginate_queryset(queryset)
        
        if page is not None:
            serializer = self.get_serializer(page, many=True)
            paginated_response = self.get_paginated_response(serializer.data)
            
            # Custom response format
            return Response({
                "data": serializer.data,
                "meta": {
                    "total_count": paginated_response.data['count'],
                    "page": request.query_params.get('page', 1),
                    "per_page": request.query_params.get('page_size', 20)
                }
            })
        
        serializer = self.get_serializer(queryset, many=True)
        return Response({
            "data": serializer.data,
            "meta": {
                "total_count": queryset.count(),
                "page": 1,
                "per_page": queryset.count()
            }
        })
    
    def get_queryset(self):
        """
        Filter journals by patient if user is a patient
        """
        # Return base queryset for swagger schema generation
        if getattr(self, 'swagger_fake_view', False):
            return PatientJournal.objects.none()

        queryset = PatientJournal.objects.all().order_by('-created_at')

        # If user is a patient, only show their journals
        if hasattr(self.request.user, 'roles') and self.request.user.roles.filter(name='patient').exists():
            try:
                patient = Patient.objects.get(user=self.request.user)
                queryset = queryset.filter(patient=patient)
            except Patient.DoesNotExist:
                queryset = PatientJournal.objects.none()
        
        # Filter by patient_id if provided (for doctors/admins)
        patient_id = self.request.query_params.get('patient_id')
        if patient_id and hasattr(self.request.user, 'roles') and self.request.user.roles.filter(name__in=['doctor', 'admin']).exists():
            queryset = queryset.filter(patient_id=patient_id)
        
        # Filter by tags
        tags = self.request.query_params.get('tags')
        if tags:
            tag_list = [tag.strip() for tag in tags.split(',')]
            for tag in tag_list:
                queryset = queryset.filter(tags__contains=[tag])
        
        return queryset
    
    def perform_create(self, serializer):
        """
        Automatically set the patient when creating a journal entry
        """
        if self.request.user.roles.filter(name='patient').exists():
            try:
                patient = Patient.objects.get(user=self.request.user)
                serializer.save(patient=patient)
            except Patient.DoesNotExist:
                raise serializers.ValidationError("User is not associated with a patient profile")
        else:
            # For non-patients, patient must be provided
            if not serializer.validated_data.get('patient'):
                raise serializers.ValidationError("Patient field is required")
            serializer.save()
    
    def get_permissions(self):
        """
        Patients can create and update their own journals
        Doctors and admins can view all journals
        """
        if self.action in ['create']:
            permission_classes = [IsPatientUser]
        elif self.action in ['update', 'partial_update', 'destroy']:
            permission_classes = [permissions.IsAuthenticated]
        else:
            permission_classes = [permissions.IsAuthenticated]
        return [permission() for permission in permission_classes]
    
    def perform_update(self, serializer):
        """
        Ensure patients can only update their own journals
        """
        journal = self.get_object()
        if (self.request.user.roles.filter(name='patient').exists() and 
            journal.patient.user != self.request.user):
            raise serializers.ValidationError("You can only edit your own journal entries")
        serializer.save()
    
    def perform_destroy(self, instance):
        """
        Ensure patients can only delete their own journals
        """
        if (self.request.user.roles.filter(name='patient').exists() and 
            instance.patient.user != self.request.user):
            raise serializers.ValidationError("You can only delete your own journal entries")
        instance.delete()
